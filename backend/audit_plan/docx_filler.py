"""
BATUHAN — Audit Plan: DOCX Filler
Injects the generated schedule into Table 2 of the pre-filled FR.223 template.

Strategy:
  1. Open uploaded template bytes with python-docx.
  2. Locate Table 2 (schedule table, index 2).
  3. Keep header row (Row 0), delete all subsequent rows.
  4. For each day: insert a day-header row, then one row per slot.
     Break slots get a single merged row.
  5. Save document to bytes and return.

Column widths (DXA units, total = 9883):
  Hour = 1728, Standard = 1280, Clause No = 1414, Activity = 3705, Team = 1756
"""

from __future__ import annotations
from io import BytesIO
from typing import TYPE_CHECKING
from lxml import etree

from docx import Document
from docx.oxml.ns import qn

from .schedule_generator import DaySchedule, Slot

if TYPE_CHECKING:
    from .template_reader import AuditPlanContext

# ---------------------------------------------------------------------------
# Column widths in DXA (1/20th of a point)
# ---------------------------------------------------------------------------
_COL_WIDTHS = [1728, 1280, 1414, 3705, 1756]
_TABLE_WIDTH = sum(_COL_WIDTHS)   # 9883
_FONT        = "Times New Roman"
_FONT_SIZE   = "20"               # half-points → 10 pt


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _w(tag: str) -> str:
    """Return a fully-qualified OOXML tag, e.g. 'w:tr'."""
    return f"{{{qn('w:tr').split('}')[0][1:]}}}{tag}"


def _make_rpr(bold: bool = False) -> etree._Element:
    rpr = etree.SubElement(etree.Element(_w("rPr")), _w("rPr"))
    rpr = etree.Element(qn("w:rPr"))
    fonts = etree.SubElement(rpr, qn("w:rFonts"))
    fonts.set(qn("w:ascii"), _FONT)
    fonts.set(qn("w:hAnsi"), _FONT)
    if bold:
        etree.SubElement(rpr, qn("w:b"))
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), _FONT_SIZE)
    szCs = etree.SubElement(rpr, qn("w:szCs"))
    szCs.set(qn("w:val"), _FONT_SIZE)
    return rpr


def _make_para(text: str, bold: bool = False, center: bool = False) -> etree._Element:
    p = etree.Element(qn("w:p"))
    ppr = etree.SubElement(p, qn("w:pPr"))
    if center:
        jc = etree.SubElement(ppr, qn("w:jc"))
        jc.set(qn("w:val"), "center")
    prpr = etree.SubElement(ppr, qn("w:rPr"))
    fonts = etree.SubElement(prpr, qn("w:rFonts"))
    fonts.set(qn("w:ascii"), _FONT); fonts.set(qn("w:hAnsi"), _FONT)
    if bold:
        etree.SubElement(prpr, qn("w:b"))
    sz = etree.SubElement(prpr, qn("w:sz")); sz.set(qn("w:val"), _FONT_SIZE)
    if text:
        r = etree.SubElement(p, qn("w:r"))
        r.append(_make_rpr(bold))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return p


def _make_tc(width: int, text: str, bold: bool = False,
             center: bool = False, grid_span: int = 1) -> etree._Element:
    tc = etree.Element(qn("w:tc"))
    tcpr = etree.SubElement(tc, qn("w:tcPr"))
    tcw = etree.SubElement(tcpr, qn("w:tcW"))
    tcw.set(qn("w:w"), str(width))
    tcw.set(qn("w:type"), "dxa")
    if grid_span > 1:
        gs = etree.SubElement(tcpr, qn("w:gridSpan"))
        gs.set(qn("w:val"), str(grid_span))
    vAlign = etree.SubElement(tcpr, qn("w:vAlign"))
    vAlign.set(qn("w:val"), "center")
    tc.append(_make_para(text, bold=bold, center=center))
    return tc


def _make_trpr(height: int = 397) -> etree._Element:
    trpr = etree.Element(qn("w:trPr"))
    etree.SubElement(trpr, qn("w:cantSplit"))
    trh = etree.SubElement(trpr, qn("w:trHeight"))
    trh.set(qn("w:val"), str(height))
    return trpr


# ---------------------------------------------------------------------------
# Row builders
# ---------------------------------------------------------------------------

def _unique_row_cells(row) -> list:
    """Return deduplicated cell objects for a row (handles merged cells)."""
    seen: set[int] = set()
    cells = []
    for cell in row.cells:
        if id(cell) not in seen:
            seen.add(id(cell))
            cells.append(cell)
    return cells


def _write_cell(cell, text: str) -> None:
    """Clear a table cell and write plain text, preserving the cell element itself."""
    tc = cell._tc
    # Remove all existing paragraphs
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    # Add a single paragraph with the new text
    p = etree.SubElement(tc, qn("w:p"))
    if text:
        r = etree.SubElement(p, qn("w:r"))
        t = etree.SubElement(r, qn("w:t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _fill_sites(doc: Document, ctx: "AuditPlanContext") -> None:
    """
    Populate the three site data rows in Table 1 from ctx.
    If no additional sites are defined, fills row 1 with the HQ address/scope/employees.
    Table 1 layout (fixed):
      R0 — header (Site/s | Address | Process/Activity | Number of Employees)
      R1-R3 — site data rows (we write to unique cells [1], [2], [3] of each)
      R4+ — empty separator and audit team rows (leave untouched)
    """
    if len(doc.tables) < 2:
        return

    tbl = doc.tables[1]
    if len(tbl.rows) < 4:
        return  # Template doesn't have the expected structure

    # Build site records: use ctx.sites if available, otherwise HQ as site 1
    records = []
    if ctx.sites:
        for s in ctx.sites:
            records.append((s.address, s.process, s.employees))
    else:
        records.append((ctx.address, ctx.scope[:80] if ctx.scope else "", ctx.num_employees))

    # Fill rows 1, 2, 3 (up to available records)
    site_row_indices = [1, 2, 3]
    for i, rec in enumerate(records):
        if i >= len(site_row_indices):
            break
        row = tbl.rows[site_row_indices[i]]
        cells = _unique_row_cells(row)
        # cells[0] = "Site/s" label — do NOT modify (may be vertically merged)
        if len(cells) > 1:
            _write_cell(cells[1], str(rec[0]))   # Address
        if len(cells) > 2:
            _write_cell(cells[2], str(rec[1]))   # Process/Activity
        if len(cells) > 3:
            _write_cell(cells[3], str(rec[2]))   # Number of Employees


def _day_header_row(day_number: int, date: str, site: str) -> etree._Element:
    """Full-width merged row for a new day."""
    label = f"{day_number}. Day ({date})   {site}"
    tr = etree.Element(qn("w:tr"))
    tr.append(_make_trpr())
    tr.append(_make_tc(_TABLE_WIDTH, label, bold=True, grid_span=5))
    return tr


def _data_row(slot: Slot) -> etree._Element:
    """Standard 5-cell schedule row."""
    tr = etree.Element(qn("w:tr"))
    tr.append(_make_trpr())
    tr.append(_make_tc(_COL_WIDTHS[0], slot.time,      center=True))
    tr.append(_make_tc(_COL_WIDTHS[1], slot.standard,  center=True))
    tr.append(_make_tc(_COL_WIDTHS[2], slot.clauses,   center=True))
    tr.append(_make_tc(_COL_WIDTHS[3], slot.activity))
    tr.append(_make_tc(_COL_WIDTHS[4], slot.auditors,  center=True))
    return tr


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def fill_schedule(
    docx_bytes: bytes,
    days: list[DaySchedule],
    ctx: "AuditPlanContext | None" = None,
) -> bytes:
    """
    Fill Table 1 (sites) and Table 2 (schedule) of the uploaded template.

    Args:
        docx_bytes: Raw bytes of the uploaded pre-filled FR.223 template.
        days:       Generated schedule (list of DaySchedule).
        ctx:        Full AuditPlanContext; when provided, Table 1 site rows are populated.

    Returns:
        Bytes of the completed .docx file ready for download.

    Raises:
        ValueError: If the template lacks the expected tables.
    """
    doc = Document(BytesIO(docx_bytes))

    if len(doc.tables) < 3:
        raise ValueError(
            f"Template has only {len(doc.tables)} table(s); expected at least 3. "
            "Ensure you uploaded the correct FR.223 audit plan template."
        )

    # ---- Table 1: fill site rows ----
    if ctx is not None:
        _fill_sites(doc, ctx)

    # ---- Table 2: fill schedule ----
    tbl = doc.tables[2]
    tbl_elem = tbl._tbl

    # Keep the header row (index 0); remove all other rows
    rows_to_remove = list(tbl_elem.findall(qn("w:tr")))[1:]
    for row in rows_to_remove:
        tbl_elem.remove(row)

    # Insert generated rows.
    # Break slots (is_break=True) are rendered as normal 5-cell rows so the
    # Activity column displays "Lunch Break" rather than a merged cell.
    for day in days:
        tbl_elem.append(_day_header_row(day.day_number, day.date, day.site))
        for slot in day.slots:
            tbl_elem.append(_data_row(slot))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
