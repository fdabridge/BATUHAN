"""
BATUHAN — Audit Plan: Template Reader
Reads a pre-filled FR.223 DOCX (Tables 0 and 1 already filled by the user)
and extracts all context needed for schedule generation.

Table layout (from real documents):
  Table 0 — header  (15 rows, 5 cols, many merged cells)
  English/UAF: sites + team share Table 1; schedule is Table 2.
  Turkish/TÜRKAK: sites and team are separate; schedule is Table 4.
  Signature tables are never modified.
"""

from __future__ import annotations
import re
import unicodedata
from dataclasses import dataclass, field
from io import BytesIO

from docx import Document

from .clause_map import normalize_standard, normalize_audit_type


# ---------------------------------------------------------------------------
# Data structures returned to the caller
# ---------------------------------------------------------------------------

@dataclass
class AuditorEntry:
    role: str           # "Lead Auditor" | "Auditor" | "Trainee Auditor" | "Technical Expert" | "Observer"
    name: str
    standard: str       # raw text from template
    ea_code: str


@dataclass
class SiteEntry:
    address: str
    process: str
    employees: str


@dataclass
class DayWindow:
    date: str
    start_time: str
    end_time: str
    lunch_start: str = "13.00"
    lunch_end: str = "14.00"
    site: str = ""


@dataclass
class AuditPlanContext:
    # Header fields
    date: str
    project_number: str
    org_name: str
    address: str
    telephone: str
    email: str
    org_representative: str
    standards_raw: str          # raw text, e.g. "ISO 9001:2015, ISO 14001:2015"
    standards: list[str]        # normalised CLAUSE_MAP keys
    ea_code: str
    category: str
    scope: str
    not_applicable: str
    audit_type_raw: str         # raw text from template
    audit_type: str             # normalised CLAUSE_MAP key, e.g. "Stage 2"
    audit_dates: str            # raw string, e.g. "27-28.09.2025"
    num_employees: str
    audit_time: str
    shift_number: str
    language: str
    audit_criteria: str
    audit_objectives: str
    # Team & sites
    auditors: list[AuditorEntry] = field(default_factory=list)
    sites: list[SiteEntry] = field(default_factory=list)
    day_windows: list[DayWindow] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _unique_cells(row) -> list[str]:
    """Return deduplicated cell texts for a row (handles merged cells)."""
    seen: set[int] = set()
    texts: list[str] = []
    for cell in row.cells:
        if id(cell) not in seen:
            seen.add(id(cell))
            texts.append(cell.text.strip())
    return texts


def _find_value(table, *labels: str) -> str:
    """
    Scan every row of *table* for a cell containing any of *labels* (case-insensitive).
    Return the text of the NEXT unique cell in the same row.
    """
    label_set = {lbl.lower() for lbl in labels}
    for row in table.rows:
        cells = _unique_cells(row)
        for i, text in enumerate(cells):
            if any(lbl in text.lower() for lbl in label_set):
                if i + 1 < len(cells) and cells[i + 1]:
                    return cells[i + 1]
    return ""


def _normalise_standards(raw: str) -> list[str]:
    """Parse full ISO names or management-system aliases from integrated text."""
    standards: list[str] = []
    for part in re.split(r"[,\n;+]+", raw):
        norm = normalize_standard(part.strip())
        if norm and norm not in standards:
            standards.append(norm)
    return standards


def _fold_label(value: str) -> str:
    value = unicodedata.normalize("NFKD", value or "")
    value = "".join(char for char in value if not unicodedata.combining(char))
    return value.lower().translate(str.maketrans({"ı": "i", "ş": "s", "ğ": "g"}))


_TEAM_ROLE_ALIASES = {
    "lead auditor": "Lead Auditor",
    "basdenetci": "Lead Auditor",
    "auditor": "Auditor",
    "denetci": "Auditor",
    "trainee auditor": "Trainee Auditor",
    "aday denetci": "Trainee Auditor",
    "technical expert": "Technical Expert",
    "technical experts": "Technical Expert",
    "teknik uzman": "Technical Expert",
    "observer": "Observer",
    "gozlemci": "Observer",
}


def _canonical_team_role(raw: str) -> str | None:
    folded = _fold_label(raw).strip()
    for alias, canonical in _TEAM_ROLE_ALIASES.items():
        if alias in folded:
            return canonical
    return None


def _is_team_header(cells: list[str]) -> bool:
    folded = _fold_label(" ".join(cells))
    return "name surname" in folded or "ad soyad" in folded


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def read_template(docx_bytes: bytes) -> AuditPlanContext:
    """
    Parse a pre-filled FR.223 DOCX and return an AuditPlanContext.

    Args:
        docx_bytes: Raw bytes of the uploaded .docx file.

    Returns:
        AuditPlanContext with all header, team and site data extracted.

    Raises:
        ValueError: If the file cannot be parsed or lacks required tables.
    """
    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError(f"Cannot open uploaded file as a DOCX: {exc}") from exc

    if len(doc.tables) < 3:
        raise ValueError(
            f"Expected at least 3 tables in the audit plan template, found {len(doc.tables)}. "
            "Please upload the correct pre-filled FR.223 blank template."
        )

    tbl0 = doc.tables[0]
    tbl1 = doc.tables[1]

    # ---- Table 0: header fields ----
    date            = _find_value(tbl0, "date")
    project_number  = _find_value(tbl0, "project no", "project number")
    org_name        = _find_value(tbl0, "organization", "organisation")
    address         = _find_value(tbl0, "address")
    telephone       = _find_value(tbl0, "telephone")
    email           = _find_value(tbl0, "e-mail", "email")
    org_rep         = _find_value(tbl0, "organisation representative", "organization representative")
    standards_raw   = _find_value(tbl0, "standard/s", "standards")
    ea_code         = _find_value(tbl0, "ea/nace", "ea/iaf")
    scope           = _find_value(tbl0, "scope")
    not_applicable  = _find_value(tbl0, "not applicable")
    audit_type_raw  = _find_value(tbl0, "audit type")
    audit_dates     = _find_value(tbl0, "audit date")
    num_employees   = _find_value(tbl0, "number of effective employees", "effective employees")
    audit_time      = _find_value(tbl0, "audit time")
    shift_number    = _find_value(tbl0, "shift number")
    language        = _find_value(tbl0, "audit language", "language")
    audit_criteria  = _find_value(tbl0, "audit criteria")
    audit_objectives= _find_value(tbl0, "audit objectives")

    # Normalise standards (may be comma/newline separated)
    standards = _normalise_standards(standards_raw)

    # Normalise audit type
    audit_type = normalize_audit_type(audit_type_raw) or "Stage 2"

    # ---- Sites + audit team (combined in English, separate in Turkish) ----
    auditors: list[AuditorEntry] = []
    sites: list[SiteEntry] = []
    SITE_HEADER_KEYWORDS = {
        "site/s", "site address", "address", "process/activity",
        "saha/lar", "adres", "proses/faaliyet",
    }

    team_tables = []
    for table in doc.tables[1:]:
        if any(_is_team_header(_unique_cells(row)) for row in table.rows):
            team_tables.append(table)

    # Sites always live in the second table, even when the Turkish team is separate.
    for row in tbl1.rows:
        cells = _unique_cells(row)
        if not cells or not cells[0]:
            continue
        first = _fold_label(cells[0]).strip()
        if _is_team_header(cells):
            break
        # Site rows: skip header row, grab data rows
        if any(kw in first for kw in SITE_HEADER_KEYWORDS):
            continue
        addr    = cells[1] if len(cells) > 1 else ""
        process = cells[2] if len(cells) > 2 else ""
        emps    = cells[3] if len(cells) > 3 else ""
        if addr:
            sites.append(SiteEntry(address=addr, process=process, employees=emps))

    for table in team_tables:
        in_team_section = False
        for row in table.rows:
            cells = _unique_cells(row)
            if _is_team_header(cells):
                in_team_section = True
                continue
            if not in_team_section or not cells:
                continue
            role = _canonical_team_role(cells[0])
            if not role:
                continue
            name = cells[1] if len(cells) > 1 else ""
            std = cells[2] if len(cells) > 2 else ""
            ea = cells[3] if len(cells) > 3 else ""
            if name:
                auditors.append(AuditorEntry(
                    role=role,
                    name=name,
                    standard=std,
                    ea_code=ea,
                ))

    return AuditPlanContext(
        date=date, project_number=project_number, org_name=org_name,
        address=address, telephone=telephone, email=email,
        org_representative=org_rep, standards_raw=standards_raw,
        standards=standards, ea_code=ea_code, category="", scope=scope,
        not_applicable=not_applicable, audit_type_raw=audit_type_raw,
        audit_type=audit_type, audit_dates=audit_dates,
        num_employees=num_employees, audit_time=audit_time,
        shift_number=shift_number, language=language,
        audit_criteria=audit_criteria, audit_objectives=audit_objectives,
        auditors=auditors, sites=sites,
    )
