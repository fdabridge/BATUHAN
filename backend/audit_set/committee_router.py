"""
BATUHAN — Certification Committee appointment (Prompt 14).

Endpoints under /audit-sets:
  GET    /audit-sets/{id}/committee
  GET    /audit-sets/{id}/committee/eligible-users
  GET    /audit-sets/{id}/committee/cert-manager
  POST   /audit-sets/{id}/committee/appoint
  DELETE /audit-sets/{id}/committee/{member_id}
  GET    /audit-sets/{id}/planning/committee/available-auditors  (Portal 64)
  POST   /audit-sets/{id}/fr233/generate
  POST   /audit-sets/{id}/fr233/release
  POST   /audit-sets/{id}/fr233/upload
  GET    /audit-sets/{id}/fr233
"""
from __future__ import annotations

from fastapi import APIRouter, Body, Depends, File, HTTPException, Query, UploadFile
from pydantic import BaseModel
from sqlalchemy import and_, or_
from sqlalchemy.orm import Session

# Portal 62 — short-code → ISO-name mapping used to match audit_set.standards
# (e.g. "QMS") against AuditorStandardQualification.standard_code ("ISO 9001").
# Mirrors backend/audit_set/service.py _CODE_TO_ISO to avoid cross-module import.
_STD_CODE_TO_ISO: dict[str, str] = {
    "QMS":        "ISO 9001",
    "EMS":        "ISO 14001",
    "OHSMS":      "ISO 45001",
    "FSMS":       "ISO 22000",
    "FSSC 22000": "FSSC 22000",
    "MDQMS":      "ISO 13485",
    "MDMS":       "ISO 13485",
    "ISMS":       "ISO 27001",
    "ENMS":       "ISO 50001",
    "EnMS":       "ISO 50001",
    "ABMS":       "ISO 37001",
    "CMS":        "ISO 37301",
}

from storage.document_store import upload as store_upload, ensure_local, delete as store_delete, invalidate_cache, is_s3_ref, resolve_docx_key
from audit_set.committee_slots import (
    committee_member_auditor_id,
    committee_member_name,
    planned_committee_slots,
)
from audit_set.db_models import (
    AuditSet, AuditSetCommitteeMember, AuditSetFR233Record, AuditSetSharedDocument,
    AuditSetStage, AuditDocumentSignature, VisualSignaturePlacement, get_db,
)
from auth.db_models import PlatformUser, get_db as get_auth_db
from auth.dependencies import get_current_user
from auth.policy import resolve_realtime_action_datetime
from auditors.qualification_scope import normalize_scope_code, normalize_scope_type

class GenerateFR233Request(BaseModel):
    released_at_override: str | None = None
    """ISO 8601 date or datetime string for a retroactive release date."""

class ReleaseFR233Request(BaseModel):
    released_at: str | None = None
    """ISO 8601 date for retroactive release. Format: YYYY-MM-DD or YYYY-MM-DDTHH:MM:SS.
    Defaults to current UTC time if omitted."""

router = APIRouter(prefix="/audit-sets", tags=["committee"])

CB_ROLES = {"admin", "planner", "planner_us", "officer", "executive", "gm", "certification_manager"}
CM_ROLES = {"admin", "executive"}   # roles that act as Certification Manager


def _collect_stage_auditor_ids(stages: list) -> set[str]:
    """Return the set of auditors.auditors.id values assigned to any stage."""
    ids: set[str] = set()
    for s in stages:
        if s.lead_auditor_id:
            ids.add(s.lead_auditor_id)
        for group in (s.auditors or [], s.technical_experts or [],
                      s.observers or [], s.ik_experts or [], s.evaluators or []):
            for p in group:
                if isinstance(p, dict) and p.get("id"):
                    ids.add(p["id"])
    return ids


def _merge_isms_application_scope(
    required_scope: dict,
    audit_standards: set[str],
    application_data: dict | None,
) -> dict:
    """Repair legacy ISO 27001 scope without replacing planner selections."""
    merged = dict(required_scope)
    app_data = application_data or {}
    isms_areas = app_data.get("isms_technical_areas") or []
    if isinstance(isms_areas, str):
        isms_areas = [area.strip() for area in isms_areas.split(",") if area.strip()]
    if not isms_areas and app_data.get("isms_technical_area"):
        isms_areas = [app_data["isms_technical_area"]]
    isms_areas = list(dict.fromkeys(isms_areas))
    if not isms_areas:
        return merged

    for iso_std in audit_standards:
        if "27001" not in iso_std:
            continue
        current = merged.get(iso_std) or {}
        if current.get("type") != "isms" or not current.get("codes"):
            merged[iso_std] = {"type": "isms", "codes": isms_areas}
    return merged


@router.get("/{audit_set_id}/committee")
def get_committee(
    audit_set_id: str,
    db: Session = Depends(get_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    if current_user.role not in CB_ROLES:
        raise HTTPException(403, "Not authorized")

    members = (
        db.query(AuditSetCommitteeMember)
        .filter_by(audit_set_id=audit_set_id)
        .order_by(AuditSetCommitteeMember.appointed_at)
        .all()
    )
    reviewer_sigs = {
        s.signer_user_id: s
        for s in db.query(AuditDocumentSignature).filter_by(
            audit_set_id=audit_set_id,
            signer_role_label="cb_reviewer",
        ).all()
        if s.signer_user_id
    }
    return [
        {
            "id":                      m.id,
            "user_id":                 m.user_id,
            "user_name":               m.user_name,
            "user_email":              m.user_email,
            "role":                    m.role,
            "appointed_by":            m.appointed_by,
            "ea_codes_at_appointment": m.ea_codes_at_appointment,
            "appointed_at":            m.appointed_at.isoformat() if m.appointed_at else None,
            "has_signed_fr218":        (
                reviewer_sigs.get(m.user_id) is not None
                and reviewer_sigs[m.user_id].signed_at is not None
            ),
        }
        for m in members
    ]


@router.get("/{audit_set_id}/committee/eligible-users")
def get_eligible_users(
    audit_set_id: str,
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    if current_user.role not in CB_ROLES:
        raise HTTPException(403, "Not authorized")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    stages = db.query(AuditSetStage).filter_by(audit_set_id=audit_set_id).all()
    stage_auditor_ids = _collect_stage_auditor_ids(stages)

    already_appointed_user_ids = {
        m.user_id for m in
        db.query(AuditSetCommitteeMember).filter_by(audit_set_id=audit_set_id).all()
    }

    # Portal 54 — eligible candidates are CB staff PLUS external auditors with
    # a linked auditor profile. Previously CB-only filtering excluded auditors
    # from the committee picker even when they covered the right EA codes.
    candidate_users = (
        auth_db.query(PlatformUser)
        .filter(
            PlatformUser.is_active == True,  # noqa: E712
            or_(
                PlatformUser.role.in_(CB_ROLES),
                and_(
                    PlatformUser.role == "auditor",
                    PlatformUser.auditor_id.isnot(None),
                ),
            ),
        )
        .all()
    )

    from auditors.models import Auditor as AuditorModel

    plan_ea_code = (audit_set.ea_code or "").strip()

    results = []
    for u in candidate_users:
        if u.id in already_appointed_user_ids:
            continue

        ea_codes: list[str] = []
        on_audit_team = False

        if u.auditor_id:
            if u.auditor_id in stage_auditor_ids:
                on_audit_team = True
            auditor = db.query(AuditorModel).filter_by(id=u.auditor_id).first()
            if auditor:
                ea_codes = auditor.ea_codes or []

        if on_audit_team:
            continue

        if ea_codes:
            ea_match = (not plan_ea_code) or (plan_ea_code in ea_codes)
        else:
            ea_match = True

        results.append({
            "user_id":              u.id,
            "full_name":            u.full_name,
            "email":                u.email,
            "role":                 u.role,
            "ea_codes":             ea_codes,
            "ea_match":             ea_match,
            "has_auditor_profile":  bool(u.auditor_id),
            "eligible_as_reviewer": ea_match,
        })

    results.sort(key=lambda x: (not x["eligible_as_reviewer"], x["full_name"]))
    return results


class AppointRequest(BaseModel):
    # Portal 61 — user_id optional for role="reviewer" (auto-resolves to the
    # system's certification_manager). Still required for decision_maker.
    user_id: str | None = None
    role: str  # "reviewer" | "decision_maker"


@router.get("/{audit_set_id}/committee/cert-manager")
def get_cert_manager(
    audit_set_id: str,
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    """Portal 61 — return the single active Certification Manager for display
    in the auto-assign reviewer confirmation UI."""
    if current_user.role not in CB_ROLES:
        raise HTTPException(403, "Not authorized")
    if not db.query(AuditSet).filter_by(id=audit_set_id).first():
        raise HTTPException(404, "Audit set not found")
    cm = auth_db.query(PlatformUser).filter_by(
        role="certification_manager", is_active=True,
    ).first()
    if not cm:
        return {"cert_manager": None}
    return {"cert_manager": {"user_id": cm.id, "full_name": cm.full_name, "email": cm.email}}


# ── Portal 62 — FR.233 certification committee team (auditor pool) ───────────

def _auditor_iso_quals(db: Session, auditor_id: str) -> set[str]:
    """Return the set of ISO standard codes the auditor is currently qualified for."""
    from auditors.models import AuditorStandardQualification
    rows = (
        db.query(AuditorStandardQualification.standard_code)
        .filter(AuditorStandardQualification.auditor_id == auditor_id)
        .filter(AuditorStandardQualification.is_qualified == True)  # noqa: E712
        .all()
    )
    return {r[0] for r in rows if r[0]}


def _audit_iso_standards(audit_set) -> set[str]:
    """Return audit_set.standards mapped to ISO names ({'ISO 9001', ...})."""
    return {_STD_CODE_TO_ISO.get(s, s) for s in (audit_set.standards or [])}


@router.get("/{audit_set_id}/planning/committee/available-auditors")
def get_planning_committee_available(
    audit_set_id: str,
    exclude_auditor_ids: str | None = Query(default=None),
    db: Session = Depends(get_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    """Portal 64 — available auditors for the committee picker during planning.

    Returns each active auditor not in `exclude_auditor_ids` with:
      - `covered_scope`: {iso_std: [matched_codes]} — same format as
        AuditorAvailabilityItem.covered_scope from GET /auditors/available
      - `covers_audit`: true when the auditor covers every required ISO standard
        AND the audit's EA code (numeric-normalised to handle "EA 5" vs "5")

    Auditors with `covers_audit=true` are sorted first.
    """
    if current_user.role not in CB_ROLES:
        raise HTTPException(403, "Not authorized")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    excluded: set[str] = set()
    if exclude_auditor_ids:
        excluded = {aid.strip() for aid in exclude_auditor_ids.split(",") if aid.strip()}

    from auditors.models import Auditor as AuditorModel

    plan_ea_code    = (audit_set.ea_code or "").strip()
    audit_standards = _audit_iso_standards(audit_set)   # {"ISO 9001", ...}

    # Build required_scope dict (same format used by /auditors/available).
    # Prefer the stored required_scope when available; fall back to deriving
    # it from standards + ea_code for older audit sets.
    if audit_set.required_scope:
        req_cat: dict = dict(audit_set.required_scope)
    elif audit_standards and plan_ea_code:
        req_cat = {iso_std: {"type": "ea", "codes": [plan_ea_code]}
                   for iso_std in audit_standards}
    else:
        req_cat = {}

    # Existing audit sets may still carry the old ISO 27001 EA-code scope.
    # Prefer the ISMS technical area captured on the application.
    req_cat = _merge_isms_application_scope(
        req_cat,
        audit_standards,
        audit_set.application_data,
    )

    def _ea_int(code) -> int | None:
        """Strip 'EA' prefix and return the integer sector number, or None.
        Handles integer inputs (e.g. 12) in addition to string inputs
        ("EA 12", "12") so that JSON fields storing codes as bare numbers
        still compare correctly.
        """
        try:
            if isinstance(code, (int, float)):
                return int(code)
            return int(str(code).strip().upper().replace("EA", "").replace(" ", ""))
        except (ValueError, AttributeError):
            return None

    target_ea = _ea_int(plan_ea_code) if plan_ea_code else None

    def _compute_covered_scope(qualifications: list, req: dict, top_level_ea: list) -> dict:
        """For each required ISO standard + required codes, check whether this
        auditor's qualification record covers any of those codes.
        Returns {iso_std: [matched_codes]}.

        EA codes are compared numerically ("EA 36" == "EA36" == "36") so that
        format differences between the audit set and the auditor profile do not
        silently drop auditors from the committee picker.
        """
        covered: dict = {}
        for iso_std, entry in req.items():
            scope_type      = normalize_scope_type(entry.get("type", "ea"))
            required_codes: list = entry.get("codes", []) or []
            if not required_codes:
                continue
            std_lower = iso_std.lower().replace("iso ", "").replace(" ", "")
            qual = next(
                (q for q in (qualifications or [])
                 if q.is_qualified is not False and std_lower in
                 (q.standard_code or "").lower().replace("iso ", "").replace(" ", "")),
                None,
            )
            if scope_type in ("food", "medical", "isms", "sector", "energy"):
                if not qual:
                    continue
                raw = qual.scope_category or ""
                auditor_codes = [c.strip() for c in raw.split(",") if c.strip()]
                auditor_keys = {normalize_scope_code(c, scope_type) for c in auditor_codes}
                matched = [
                    c for c in required_codes
                    if normalize_scope_code(c, scope_type) in auditor_keys
                ]
            else:
                # EA codes — numeric normalisation: "EA 36" == "EA36" == "36".
                # Committee review is sector-competence based: if the auditor
                # covers the required EA code for any EA-standard qualification
                # or in their top-level EA list, that EA sector can cover the
                # audit's EA requirement across ISO 9001/14001/45001.
                auditor_codes: list = []
                for q in (qualifications or []):
                    if q.is_qualified is not False:
                        auditor_codes.extend(q.ea_codes or [])
                if not auditor_codes:
                    auditor_codes = top_level_ea
                aud_ints = {_ea_int(c) for c in auditor_codes} - {None}
                matched  = [c for c in required_codes if _ea_int(c) in aud_ints]
            if matched:
                covered[iso_std] = matched
        return covered

    auditors = (
        db.query(AuditorModel)
        .filter(AuditorModel.is_active == True)  # noqa: E712
        .all()
    )

    results = []
    for a in auditors:
        if a.id in excluded:
            continue

        ea_codes = a.ea_codes or []

        # EA coverage — check per-standard EA codes first, fall back to top-level.
        # Use numeric normalisation so "EA 5" matches "5" and vice-versa.
        if plan_ea_code and target_ea is not None:
            per_std_ea: list[str] = []
            for q in (a.standard_qualifications or []):
                if q.is_qualified is not False:
                    per_std_ea.extend(q.ea_codes or [])
            codes_to_check = per_std_ea if per_std_ea else ea_codes
            ea_ok = any(_ea_int(c) == target_ea for c in codes_to_check)
        else:
            ea_ok = True

        iso_q  = _auditor_iso_quals(db, a.id)
        std_ok = (not audit_standards) or bool(audit_standards & iso_q)

        covered_scope = _compute_covered_scope(a.standard_qualifications, req_cat, a.ea_codes or [])

        results.append({
            "id":            a.id,
            "full_name":     a.name,
            "email":         a.email,
            "ea_codes":      ea_codes,
            "standards":     sorted(iso_q),
            "covers_audit":  ea_ok and std_ok,
            "covered_scope": covered_scope,
        })

    results.sort(key=lambda x: (not x["covers_audit"], x["full_name"] or ""))
    return results


@router.post("/{audit_set_id}/committee/appoint")
def appoint_committee_member(
    audit_set_id: str,
    body: AppointRequest,
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    if current_user.role not in {"admin", "planner", "planner_us"}:
        raise HTTPException(403, "Only admin or planner can appoint committee members")

    if body.role not in ("reviewer", "decision_maker"):
        raise HTTPException(400, "role must be 'reviewer' or 'decision_maker'")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    # Portal 61 — reviewer role is auto-assigned to the system's single
    # Certification Manager; any caller-supplied user_id is ignored.
    if body.role == "reviewer":
        user = auth_db.query(PlatformUser).filter_by(
            role="certification_manager", is_active=True,
        ).first()
        if not user:
            raise HTTPException(
                400, "No active Certification Manager account found",
            )
    else:
        if not body.user_id:
            raise HTTPException(400, "user_id is required for decision_maker")
        user = auth_db.query(PlatformUser).filter_by(id=body.user_id).first()
        if not user or user.role not in CB_ROLES:
            raise HTTPException(400, "User not found or not a CB user")

    already = db.query(AuditSetCommitteeMember).filter_by(
        audit_set_id=audit_set_id, user_id=user.id
    ).first()
    if already:
        raise HTTPException(409, "User is already a committee member for this audit set")

    from auditors.models import Auditor as AuditorModel
    ea_codes_snapshot: list[str] = []
    if user.auditor_id:
        auditor = db.query(AuditorModel).filter_by(id=user.auditor_id).first()
        if auditor:
            ea_codes_snapshot = auditor.ea_codes or []

    member = AuditSetCommitteeMember(
        audit_set_id=audit_set_id,
        user_id=user.id,
        user_name=user.full_name,
        user_email=user.email,
        role=body.role,
        appointed_by=current_user.id,
        ea_codes_at_appointment=ea_codes_snapshot,
    )
    db.add(member)

    if body.role == "reviewer":
        sig = (
            db.query(AuditDocumentSignature)
            .filter_by(
                audit_set_id=audit_set_id,
                document_type="FR218",
                signer_role_label="cb_reviewer",
            )
            .first()
        )
        if sig and not sig.signed_at:
            sig.signer_user_id = user.id
            sig.signer_name    = user.full_name
            sig.signer_email   = user.email

    db.commit()
    return {
        "appointed": True,
        "user_id":   user.id,
        "user_name": user.full_name,
        "role":      body.role,
    }


@router.delete("/{audit_set_id}/committee/{member_id}")
def remove_committee_member(
    audit_set_id: str,
    member_id: str,
    db: Session = Depends(get_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    if current_user.role not in {"admin", "planner", "planner_us"}:
        raise HTTPException(403, "Not authorized")

    member = db.query(AuditSetCommitteeMember).filter_by(
        id=member_id, audit_set_id=audit_set_id
    ).first()
    if not member:
        raise HTTPException(404, "Committee member not found")

    has_signed = (
        db.query(AuditDocumentSignature)
        .filter_by(audit_set_id=audit_set_id, signer_user_id=member.user_id)
        .filter(AuditDocumentSignature.signed_at.isnot(None))
        .count()
    ) > 0
    if has_signed:
        raise HTTPException(
            409, "Cannot remove a committee member who has already signed documents"
        )

    if member.role == "reviewer":
        sig = (
            db.query(AuditDocumentSignature)
            .filter_by(
                audit_set_id=audit_set_id,
                document_type="FR218",
                signer_role_label="cb_reviewer",
                signer_user_id=member.user_id,
            )
            .first()
        )
        if sig:
            sig.signer_user_id = None
            sig.signer_name    = None
            sig.signer_email   = None

    db.delete(member)
    db.commit()
    return {"removed": True}


# ── Portal 49a Part 3 — FR.233 Review & Decision Form ─────────────────────────

@router.post("/{audit_set_id}/fr233/generate")
def generate_fr233(
    audit_set_id: str,
    body: GenerateFR233Request = Body(default=GenerateFR233Request()),
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    """Render FR.233 from the blank template, persist it as a SharedDocument and
    upsert the AuditSetFR233Record. Idempotent: re-running overwrites the file
    and resets status to ``signing``."""
    import os
    from datetime import datetime
    from audit_set.fr233_generator import render_fr233_bytes
    import datetime as _dt

    if current_user.role not in {"admin", "planner", "planner_us", "executive", "certification_manager"}:
        raise HTTPException(403, "Only Planner or Certification Manager may generate FR.233")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    allowed_statuses = {
        "stage2_complete", "committee_review", "under_review",
        "stage2_in_progress",
    }
    if audit_set.workflow_status not in allowed_statuses:
        raise HTTPException(
            400,
            f"FR.233 cannot be generated while workflow_status='{audit_set.workflow_status}'. "
            "Complete Stage 2 first.",
        )

    from audit_set.workflow_router import _assert_nc_stage_complete_gate, _fr233_nc_gate_stage
    _assert_nc_stage_complete_gate(db, audit_set_id, _fr233_nc_gate_stage(db, audit_set_id))

    # Query record early so we can use it for released_at resolution.
    record = db.query(AuditSetFR233Record).filter_by(audit_set_id=audit_set_id).first()

    # Resolve released_at:
    # 1. Explicit override from caller (retroactive date).
    # 2. Preserve existing released_at from a prior blank release.
    # 3. Fall back to current time.
    _existing_doc_for_date = None
    if record and record.document_id:
        _existing_doc_for_date = db.query(AuditSetSharedDocument).filter_by(
            id=record.document_id
        ).first()

    requested_released_at = None
    if body.released_at_override:
        try:
            requested_released_at = _dt.datetime.fromisoformat(body.released_at_override)
        except ValueError:
            raise HTTPException(
                400,
                "released_at_override must be ISO 8601: YYYY-MM-DD or YYYY-MM-DDTHH:MM:SS",
            )
    elif _existing_doc_for_date and _existing_doc_for_date.released_at:
        requested_released_at = _existing_doc_for_date.released_at
    released_at_dt = resolve_realtime_action_datetime(auth_db, requested_released_at)

    # Portal 125 — allow generation even with empty committee (blank rows will render).
    committee_members = audit_set.committee_members or []
    if not committee_members:
        import logging as _log
        _log.getLogger(__name__).info(
            "generate_fr233 called with no committee members for audit_set_id=%s — "
            "rendering blank committee rows", audit_set_id,
        )

    try:
        docx_bytes = render_fr233_bytes(audit_set, db)
    except Exception as exc:
        raise HTTPException(500, f"FR.233 render failed: {exc}")

    fname = f"FR233_{audit_set.plan_number or audit_set_id}.docx"
    relative_path = f"shared_docs/{audit_set_id}/{fname}"
    out_path = store_upload(relative_path, docx_bytes)
    doc = None
    if record and record.document_id:
        doc = db.query(AuditSetSharedDocument).filter_by(id=record.document_id).first()
    if doc is None:
        doc = AuditSetSharedDocument(
            audit_set_id=audit_set_id,
            label=f"FR.233 Review & Decision — {audit_set.plan_number or ''}".strip(" —"),
            document_type="fr233",
            file_path=out_path,
            direction="cb_to_client",
            status="released",
            released_by=current_user.id,
            released_at=released_at_dt,
        )
        db.add(doc)
        db.flush()
    else:
        doc.file_path  = out_path
        doc.status     = "released"
        doc.released_by= current_user.id
        doc.released_at= released_at_dt
        # Old PDF + extracted SIG fields are stale — drop them so the viewer
        # re-converts and re-extracts from the new DOCX on next open.
        invalidate_cache(out_path)
        from audit_set.db_models import DocumentSignatureField
        _dsf_key = resolve_docx_key(out_path)
        db.query(DocumentSignatureField).filter_by(docx_path=_dsf_key).delete()

    if record is None:
        record = AuditSetFR233Record(
            audit_set_id=audit_set_id, document_id=doc.id, status="signing",
        )
        db.add(record)
    else:
        record.document_id = doc.id
        record.status      = "signing"

    if audit_set.workflow_status in {"stage2_complete", "stage2_in_progress"}:
        old = audit_set.workflow_status
        audit_set.workflow_status = "committee_review"
        from audit_set.db_models import AuditSetStatusEvent
        db.add(AuditSetStatusEvent(
            audit_set_id=audit_set_id, from_status=old, to_status="committee_review",
            triggered_by=current_user.id, notes="FR.233 generated; committee review opened",
        ))

    db.commit()
    return {
        "generated":    True,
        "document_id":  doc.id,
        "fr233_status": record.status,
    }


@router.post("/{audit_set_id}/fr233/release")
def release_fr233_blank(
    audit_set_id: str,
    body: ReleaseFR233Request = Body(default=ReleaseFR233Request()),
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    """Portal 125 — Release FR.233 as a blank document from the template."""
    import datetime as _dt
    from audit_set.fr233_generator import (
        _resolve_fr233_template,
        _safe_fill_table0,
        _fill_table3_committee,
        _build_committee_context,
    )

    if current_user.role not in {"admin", "planner", "planner_us", "executive", "certification_manager"}:
        raise HTTPException(403, "Only Planner or Certification Manager may release FR.233")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    from audit_set.workflow_router import _assert_nc_stage_complete_gate, _fr233_nc_gate_stage
    _assert_nc_stage_complete_gate(db, audit_set_id, _fr233_nc_gate_stage(db, audit_set_id))

    existing_record = db.query(AuditSetFR233Record).filter_by(audit_set_id=audit_set_id).first()
    existing_doc = None
    if existing_record and existing_record.document_id:
        existing_doc = db.query(AuditSetSharedDocument).filter_by(
            id=existing_record.document_id
        ).first()

    requested_released_at = None
    if body.released_at:
        try:
            requested_released_at = _dt.datetime.fromisoformat(body.released_at)
        except ValueError:
            raise HTTPException(400, "released_at must be ISO 8601: YYYY-MM-DD or YYYY-MM-DDTHH:MM:SS")
    elif existing_doc and existing_doc.released_at:
        requested_released_at = existing_doc.released_at
    released_at_dt = resolve_realtime_action_datetime(auth_db, requested_released_at)

    template_path = _resolve_fr233_template(audit_set)
    if not template_path:
        raise HTTPException(500, "FR.233 template not found for this audit set.")

    try:
        from docx import Document as DocxDocument
        from io import BytesIO

        docx_doc = DocxDocument(str(template_path))
        stages = {s.stage_type: s for s in (audit_set.stages or [])}
        stage1 = stages.get("stage_1")
        stage2 = stages.get("stage_2")
        auditors = [p for p in (audit_set.personnel or {}).get("auditors", []) if p.get("name")]
        team_str = ", ".join(
            f"{a['name']} (Lead Auditor)" if a.get("is_lead") else a["name"]
            for a in auditors
        )

        if len(docx_doc.tables) >= 1:
            _safe_fill_table0(docx_doc.tables[0], audit_set, team_str, stage1, stage2)
        if len(docx_doc.tables) >= 4:
            members_ctx = _build_committee_context(audit_set)
            _fill_table3_committee(docx_doc.tables[3], members_ctx)

        buf = BytesIO()
        docx_doc.save(buf)
        docx_bytes = buf.getvalue()
    except Exception as exc:
        raise HTTPException(500, f"FR.233 render failed: {exc}")

    fname = f"FR233_{audit_set.plan_number or audit_set_id}.docx"
    relative_path = f"shared_docs/{audit_set_id}/{fname}"
    out_path = store_upload(relative_path, docx_bytes)

    if existing_doc is None:
        doc_record = AuditSetSharedDocument(
            audit_set_id=audit_set_id,
            label=f"FR.233 Review & Decision — {audit_set.plan_number or ''}".strip(" —"),
            document_type="fr233",
            file_path=out_path,
            direction="cb_to_client",
            status="released",
            released_by=current_user.id,
            released_at=released_at_dt,
        )
        db.add(doc_record)
        db.flush()
    else:
        doc_record = existing_doc
        doc_record.file_path   = out_path
        doc_record.status      = "released"
        doc_record.released_by = current_user.id
        doc_record.released_at = released_at_dt
        invalidate_cache(out_path)
        from audit_set.db_models import DocumentSignatureField
        from storage.document_store import resolve_docx_key
        _dsf_key = resolve_docx_key(out_path)
        db.query(DocumentSignatureField).filter_by(docx_path=_dsf_key).delete()

    if existing_record is None:
        fr233_record = AuditSetFR233Record(
            audit_set_id=audit_set_id,
            document_id=doc_record.id,
            status="signing",
        )
        db.add(fr233_record)
    else:
        existing_record.document_id = doc_record.id
        existing_record.status      = "signing"

    if audit_set.workflow_status in {"stage2_complete", "stage2_in_progress"}:
        old = audit_set.workflow_status
        audit_set.workflow_status = "committee_review"
        from audit_set.db_models import AuditSetStatusEvent
        db.add(AuditSetStatusEvent(
            audit_set_id=audit_set_id,
            from_status=old,
            to_status="committee_review",
            triggered_by=current_user.id,
            notes="FR.233 released; committee review opened",
        ))

    db.commit()
    return {
        "released":     True,
        "document_id":  doc_record.id,
        "released_at":  released_at_dt.isoformat(),
        "fr233_status": "signing",
    }


@router.post("/{audit_set_id}/fr233/upload")
async def upload_fr233(
    audit_set_id: str,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    """Portal 57 — committee uploads the completed FR.233 Review & Decision Form
    (PDF or DOCX) prepared offline. Stored as a SharedDocument and tracked via
    AuditSetFR233Record; committee + CM sign it afterwards through the viewer.
    Mirrors the persistence path of /fr233/generate so the existing signing
    flow keeps working unchanged."""
    import os
    from datetime import datetime

    if current_user.role not in {"admin", "planner", "planner_us", "executive", "certification_manager"}:
        raise HTTPException(403, "Not authorized to upload FR.233")

    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")

    from audit_set.workflow_router import _assert_nc_stage_complete_gate, _fr233_nc_gate_stage
    _assert_nc_stage_complete_gate(db, audit_set_id, _fr233_nc_gate_stage(db, audit_set_id))

    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in {".pdf", ".docx"}:
        raise HTTPException(400, "Only PDF and DOCX files are accepted for FR.233")

    fname = f"FR233_{audit_set.plan_number or audit_set_id}{ext}"
    relative_path = f"shared_docs/{audit_set_id}/{fname}"
    contents = await file.read()
    out_path = store_upload(relative_path, contents)

    record = db.query(AuditSetFR233Record).filter_by(audit_set_id=audit_set_id).first()
    doc = None
    if record and record.document_id:
        doc = db.query(AuditSetSharedDocument).filter_by(id=record.document_id).first()
    if doc is None:
        doc = AuditSetSharedDocument(
            audit_set_id=audit_set_id,
            label=f"FR.233 Review & Decision — {audit_set.plan_number or ''}".strip(" —"),
            document_type="fr233",
            file_path=out_path,
            direction="cb_to_client",
            status="released",
            released_by=current_user.id,
            released_at=resolve_realtime_action_datetime(auth_db, None),
        )
        db.add(doc)
        db.flush()
    else:
        # Replacing the file — drop any stale rendered PDF + extracted sig fields
        # so the viewer re-converts and re-extracts on next open.
        if doc.file_path and doc.file_path != out_path:
            store_delete(doc.file_path)
        invalidate_cache(out_path)
        doc.file_path   = out_path
        doc.status      = "released"
        doc.released_by = current_user.id
        doc.released_at = resolve_realtime_action_datetime(auth_db, None)
        from audit_set.db_models import DocumentSignatureField
        _dsf_key = resolve_docx_key(out_path)
        db.query(DocumentSignatureField).filter_by(docx_path=_dsf_key).delete()

    if record is None:
        record = AuditSetFR233Record(
            audit_set_id=audit_set_id, document_id=doc.id, status="signing",
        )
        db.add(record)
    else:
        record.document_id = doc.id
        record.status      = "signing"

    if audit_set.workflow_status in {"stage2_complete", "stage2_in_progress"}:
        old = audit_set.workflow_status
        audit_set.workflow_status = "committee_review"
        from audit_set.db_models import AuditSetStatusEvent
        db.add(AuditSetStatusEvent(
            audit_set_id=audit_set_id, from_status=old, to_status="committee_review",
            triggered_by=current_user.id, notes="FR.233 uploaded; committee review opened",
        ))

    db.commit()
    return {
        "uploaded":     True,
        "document_id":  doc.id,
        "fr233_status": record.status,
    }


@router.get("/{audit_set_id}/fr233")
def get_fr233_status(
    audit_set_id: str,
    db: Session = Depends(get_db),
    auth_db: Session = Depends(get_auth_db),
    current_user: PlatformUser = Depends(get_current_user),
):
    if current_user.role not in CB_ROLES:
        raise HTTPException(403, "Not authorized")

    record = db.query(AuditSetFR233Record).filter_by(audit_set_id=audit_set_id).first()
    audit_set = db.query(AuditSet).filter_by(id=audit_set_id).first()
    if not audit_set:
        raise HTTPException(404, "Audit set not found")
    slots = planned_committee_slots(audit_set)

    # Per-slot signed lookup from VisualSignaturePlacement (source of truth).
    placements = []
    if record and record.document_id:
        placements = (
            db.query(VisualSignaturePlacement)
            .filter_by(document_type="shared_doc", doc_id=record.document_id)
            .filter(VisualSignaturePlacement.signed_at.isnot(None))
            .all()
        )
    signed_keys = {p.sig_key for p in placements}

    member_rows = []
    for static_key, member in slots.items():
        auditor_id = committee_member_auditor_id(member)
        dynamic_key = f"COMMITTEE_MEMBER_{auditor_id}" if auditor_id else None
        user = (
            auth_db.query(PlatformUser).filter_by(auditor_id=auditor_id).first()
            if auditor_id else None
        )
        is_signed = static_key in signed_keys or (
            dynamic_key is not None and dynamic_key in signed_keys
        )
        is_chair = static_key == "COMMITTEE_CHAIR"
        member_rows.append({
            "id":        auditor_id or static_key,
            "user_id":   user.id if user else "",
            "user_name": committee_member_name(member) or "",
            "role":      "decision_maker" if is_chair else "reviewer",
            "sig_key":   static_key,
            "ea_codes":  member.get("ea_codes") or [],
            "signed":    is_signed,
        })

    all_committee_signed = bool(member_rows) and all(r["signed"] for r in member_rows)

    return {
        "status":              record.status if record else "pending",
        "document_id":         record.document_id if record else None,
        "members":             member_rows,
        "cert_manager_signed": bool(
            signed_keys
            & {"CB_CERT_MANAGER", "CERT_MANAGER_FR233", "CERT_MANAGER_REVIEW"}
        ),
        "all_committee_signed": all_committee_signed,
    }
