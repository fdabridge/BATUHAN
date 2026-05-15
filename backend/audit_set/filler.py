"""
BATUHAN — Audit Set: DOCX template filler.

Opens an IFC blank-set DOCX template and writes values into the empty cells
identified by a coordinate map from `field_maps.py`.  Returns the filled
document as bytes.

The blank IFC forms have NO placeholder strings — the filler relies on
(table_idx, row_idx, col_idx) coordinates and preserves existing paragraph
and run formatting in each target cell.
"""
from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from docx import Document


from config.settings import get_settings
BLANK_SET_PATH = Path(get_settings().blank_set_path)


def fill_document(template_path: Path, field_map: dict, values: dict) -> bytes:
    """
    Open a blank DOCX template, write values into cells defined by
    `field_map`, return the filled document as bytes.

    field_map: {field_name: (table_idx, row_idx, col_idx)}
    values:    {field_name: str_value}  — only fields present in `values`
               and with a non-None value are written.
    """
    doc = Document(str(template_path))

    for field_name, (t_idx, r_idx, c_idx) in field_map.items():
        value = values.get(field_name)
        if value is None or value == "":
            continue
        try:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
        except (IndexError, AttributeError):
            continue  # Skip if coordinate doesn't exist in this variant

        # Preserve existing paragraph formatting — clear text, keep style.
        if cell.paragraphs:
            para = cell.paragraphs[0]
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = str(value)
            else:
                para.add_run(str(value))
        else:
            cell.add_paragraph(str(value))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_values(audit_set, stage=None) -> dict:
    """
    Build the values dict from an AuditSet ORM object + optional
    AuditSetStage.  All values are coerced to strings ready for cell writing.
    """
    standard_display = {
        "QMS":   "ISO 9001:2015",
        "EMS":   "ISO 14001:2015",
        "OHSMS": "ISO 45001:2018",
        "FSMS":  "ISO 22000:2018",
        "ISMS":  "ISO/IEC 27001:2022",
        "MDQMS": "ISO 13485:2016",
        "ABMS":  "ISO 37001:2016",
        "ENMS":  "ISO 50001:2018",
    }
    audit_type_display = {
        "initial":         "Initial Certification",
        "surveillance":    "Surveillance",
        "recertification": "Recertification",
    }

    standards_str = ", ".join(
        standard_display.get(s, s) for s in (audit_set.standards or [])
    )

    values: dict = {
        "plan_number":         str(audit_set.plan_number),
        "today_date":          date.today().strftime("%d/%m/%Y"),
        "company_name":        audit_set.company_name or "",
        "company_address":     audit_set.company_address or "",
        "phone":               audit_set.phone or "",
        "email":               audit_set.email or "",
        "representative":      audit_set.representative or "",
        "standards_str":       standards_str,
        "ea_code":             audit_set.ea_code or "",
        "ea_category":         audit_set.ea_category or "",
        "ea_technical_area":   audit_set.ea_technical_area or "",
        "scope_en":            audit_set.scope_en or "",
        "scope_tr":            audit_set.scope_tr or "",
        "non_applicable_clauses": audit_set.non_applicable_clauses or "",
        "audit_type_str":      audit_type_display.get(audit_set.audit_type, ""),
        "effective_employees": str(audit_set.effective_employees or ""),
        "shift_count":         str((audit_set.personnel or {}).get("shift_count", "")),
    }

    # Composite organisation block for FR.231 / FR.232 single-cell variants.
    values["organisation_block"] = "\n".join(
        x for x in [
            audit_set.company_name or "",
            audit_set.company_address or "",
            f"Tel: {audit_set.phone}" if audit_set.phone else "",
            f"E-mail: {audit_set.email}" if audit_set.email else "",
            f"Representative: {audit_set.representative}" if audit_set.representative else "",
        ] if x
    )

    # Repeated values for FR.234 second-block fields.
    values["company_name_repeat"]    = values["company_name"]
    values["company_address_repeat"] = values["company_address"]
    values["scope_repeat"]           = values["scope_en"]

    # Man-day calculation results
    mdr = audit_set.man_day_result or {}
    values["stage_1_days"]      = str(mdr.get("final_ph1", ""))
    values["stage_2_days"]      = str(mdr.get("final_ph2", ""))
    values["surveillance_days"] = str(mdr.get("final_surv1", ""))
    values["recert_days"]       = str(mdr.get("final_recert", ""))

    # Stage-specific values (auditor team, dates, audit days)
    if stage:
        def fmt_date_range(d1, d2):
            if not d1:
                return ""
            s = d1.strftime("%d/%m/%Y")
            if d2 and d2 != d1:
                s += f" – {d2.strftime('%d/%m/%Y')}"
            return s

        if stage.stage_type == "stage_1":
            values["stage_1_date"] = fmt_date_range(stage.audit_date_start, stage.audit_date_end)
        elif stage.stage_type == "stage_2":
            values["stage_2_date"] = fmt_date_range(stage.audit_date_start, stage.audit_date_end)
        elif stage.stage_type == "surveillance":
            values["surveillance_date"] = fmt_date_range(stage.audit_date_start, stage.audit_date_end)
        values["audit_days"] = str(stage.audit_days or "")

        values["lead_auditor_name"] = stage.lead_auditor_name or ""
        auditors = stage.auditors or []
        for i, a in enumerate(auditors[:3]):
            values[f"auditor_{i+1}_name"]     = a.get("name", "")
            values[f"auditor_{i+1}_standard"] = a.get("standard", "")
            values[f"auditor_{i+1}_ea"]       = a.get("ea_code", "")
        tech = stage.technical_experts or []
        for i, t in enumerate(tech[:2]):
            values[f"technical_expert_{i+1}_name"] = t.get("name", "")

    # Per-standard audit date fields for FR.222.
    for std_code in standard_display:
        s1 = values.get("stage_1_date", "")
        s2 = values.get("stage_2_date", "")
        sv = values.get("surveillance_date", "")
        values[f"stage_1_date_{std_code.lower()}"]      = s1
        values[f"stage_2_date_{std_code.lower()}"]      = s2
        values[f"surveillance_date_{std_code.lower()}"] = sv

    return values
