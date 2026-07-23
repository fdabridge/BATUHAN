"""Compatibility post-processing for FR.222 Audit Programme documents."""
from __future__ import annotations

import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


def ensure_fr222_possible_audit_dates(docx_bytes: bytes, context: dict) -> bytes:
    """Fill blank FR.222 "Possible Audit Dates" cells after rendering.

    Some still-supported blank-set copies predate the docxtpl placeholders in
    this row. They have the correct table structure, but render empty cells.
    Populate those cells from the same context used by the current template.
    Existing non-blank values are preserved.
    """
    values_by_heading = {
        "Stage 1": context.get("stage1_dates", ""),
        "Stage 2": context.get("stage2_dates", ""),
        "Surveillance 1": context.get("surv1_estimated_date", ""),
        "Surveillance 2": context.get("surv2_estimated_date", ""),
        "Recertification": context.get("recert_estimated_date", ""),
    }

    try:
        doc = Document(io.BytesIO(docx_bytes))
        changed = False

        for table in doc.tables:
            rows = table.rows
            for row_index, header_row in enumerate(rows[:-1]):
                headings = [cell.text.strip() for cell in header_row.cells]
                if (
                    "Possible Audit Dates" not in headings
                    or "Stage 1" not in headings
                    or "Stage 2" not in headings
                ):
                    continue

                date_row = rows[row_index + 1]
                date_cells = date_row.cells
                if any("Clauses to be Audited" in cell.text for cell in date_cells):
                    # This template is structurally older and has no date row.
                    # Do not risk changing its table layout during generation.
                    continue

                seen_cells: set[int] = set()
                for column, heading in enumerate(headings):
                    value = values_by_heading.get(heading)
                    if not value or column >= len(date_cells):
                        continue

                    cell = date_cells[column]
                    cell_key = id(cell._tc)
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)

                    if cell.text.strip():
                        continue

                    paragraph = cell.paragraphs[0]
                    if paragraph.runs:
                        paragraph.runs[0].text = value
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.add_run(value)
                    changed = True

        if not changed:
            return docx_bytes

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    except Exception:  # pragma: no cover - defensive: never break packaging
        logger.warning("FR.222 date-row post-process failed", exc_info=True)
        return docx_bytes
