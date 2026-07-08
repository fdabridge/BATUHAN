"""
Portal 49a Part 3 — FR.233 Review & Decision Form generator.

Renders an FR.233 DOCX for an audit set by:
  1. Resolving the correct blank template via the existing resolver.
  2. Filling project metadata (Table 0) and committee names (Table 3).
  3. Inserting ``[SIG:COMMITTEE_*]`` and ``[SIG:CERT_MANAGER_FR233]`` markers
     so the viewer's lazy field-extraction pipeline can place signatures.
"""
from __future__ import annotations

import copy
from datetime import date
from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from sqlalchemy.orm import Session

from audit_set.resolver import resolve_document_set


def _set_cell_text(tc_el, text) -> None:
    """Clear all paragraphs in a <w:tc> and write `text` into a single run on
    the first paragraph, preserving paragraph + run formatting where possible.
    Non-string values (e.g. integers stored in JSON columns) are coerced to str
    so that lxml's strict text-setter does not raise TypeError."""
    if not isinstance(text, str):
        text = "" if text is None else str(text)
    paragraphs = tc_el.findall(qn("w:p"))
    if not paragraphs:
        return
    for extra in paragraphs[1:]:
        tc_el.remove(extra)
    p = paragraphs[0]
    pPr = p.find(qn("w:pPr"))
    saved_rPr = None
    for r in p.findall(qn("w:r")):
        if saved_rPr is None:
            rPr = r.find(qn("w:rPr"))
            if rPr is not None:
                saved_rPr = copy.deepcopy(rPr)
        p.remove(r)
    if text == "":
        return
    new_r = etree.SubElement(p, qn("w:r"))
    if saved_rPr is not None:
        new_r.append(saved_rPr)
    new_t = etree.SubElement(new_r, qn("w:t"))
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")


def _fmt_d(d) -> str:
    return d.strftime("%d.%m.%Y") if d else ""


def _build_committee_context(audit_set) -> list[dict]:
    """Portal 62 — context for the docxtpl `{%tr for member in committee_members %}`
    loop. Reads the denormalized snapshot stored on AuditSet.committee_members.

    Falls back to 3 blank placeholder rows (id ``BLANK_<n>``) when the
    committee has not yet been appointed so the template still renders
    readable rows. The viewer treats `COMMITTEE_MEMBER_BLANK_*` sig_keys as
    non-signable "awaiting committee appointment" placeholders.
    """
    raw = audit_set.committee_members or []
    members = list(raw) if isinstance(raw, list) else []

    # Chairperson first.
    # Portal 124 — AuditSetCommitteeMember.role uses "decision_maker" for the chair.
    # Accept "chairperson" as legacy fallback for snapshots written before Portal 64.
    members.sort(key=lambda m: 0 if m.get("role") in ("decision_maker", "chairperson") else 1)

    ctx = [
        {
            "id":            m.get("id", ""),
            "name":          m.get("name") or m.get("full_name") or "",
            "ea_codes_str":  ", ".join(str(c) for c in (m.get("ea_codes") or [])),
            "role":          m.get("role", "member"),
        }
        for m in members
    ]

    if not ctx:
        ctx = [
            {"id": f"BLANK_{i}", "name": "", "ea_codes_str": "", "role": "member"}
            for i in range(3)
        ]

    return ctx


def _resolve_fr233_template(audit_set):
    document_set, _missing = resolve_document_set(audit_set)
    for folder, specs in document_set.items():
        for spec in specs:
            if spec.fr_number == "FR.233":
                return spec.template_path
    return None


def _fill_table3_committee(t3, members_ctx: list[dict]) -> None:
    """Fill Table 3 committee rows with member names, EA codes, and [SIG:...] markers.

    Table 3 layout (confirmed):
      Row 0: headers
      Row 1: Chairperson — col 1 = name (merged 1-2), col 3 = EA codes (merged 3-4), col 5 = sign
      Row 2: Member 1   — same column layout
      Row 3: Member 2   — same column layout
      Row 6: Certification Manager — col 4 = sign (merged 4-5)

    Sig keys assigned positionally:
      members_ctx[0] (chairperson) → COMMITTEE_CHAIR
      members_ctx[1]               → COMMITTEE_MEMBER_1
      members_ctx[2]               → COMMITTEE_MEMBER_2
    """
    _SIG_KEYS = ["COMMITTEE_CHAIR", "COMMITTEE_MEMBER_1", "COMMITTEE_MEMBER_2"]
    _COMMITTEE_ROWS = [1, 2, 3]  # row indices in Table 3

    for slot_idx, member in enumerate(members_ctx[:3]):
        row_idx = _COMMITTEE_ROWS[slot_idx]
        if row_idx >= len(t3.rows):
            break
        row = t3.rows[row_idx]
        cells = row.cells
        if len(cells) < 6:
            continue

        name = member.get("name") or ""
        ea_codes = member.get("ea_codes_str") or ""
        sig_key = _SIG_KEYS[slot_idx]

        # col 1 — name (merged with col 2; writing col 1 _tc is sufficient)
        _set_cell_text(cells[1]._tc, name)
        # col 3 — EA codes (merged with col 4; writing col 3 _tc is sufficient)
        _set_cell_text(cells[3]._tc, ea_codes)
        # col 5 — signature marker
        _set_cell_text(cells[5]._tc, f"[SIG:{sig_key}]")

    # Row 6 — Certification Manager sign cell (col 4, merged with col 5)
    if len(t3.rows) > 6:
        cm_row = t3.rows[6]
        if len(cm_row.cells) > 4:
            _set_cell_text(cm_row.cells[4]._tc, "[SIG:CERT_MANAGER_FR233]")


def render_fr233_bytes(audit_set, db: Session) -> bytes:
    """Render FR.233 bytes.

    Single-pass python-docx render:
      - Table 0: project metadata (plan number, company, dates, audit team).
      - Table 3: committee member names, EA codes, and [SIG:...] markers so the
        viewer's lazy field-extraction pipeline can place signatures.

    The docxtpl pass has been removed: the template has no Jinja2 tags.
    """
    template_path = _resolve_fr233_template(audit_set)
    if template_path is None:
        raise RuntimeError("FR.233 template not found for this audit set")

    doc = Document(str(template_path))

    stages = {s.stage_type: s for s in (audit_set.stages or [])}
    stage1 = stages.get("stage_1")
    stage2 = stages.get("stage_2")
    auditors = [p for p in (audit_set.personnel or {}).get("auditors", []) if p.get("name")]
    team_str = ", ".join(
        f"{a['name']} (Lead Auditor)" if a.get("is_lead") else a["name"]
        for a in auditors
    )

    if len(doc.tables) >= 1:
        _safe_fill_table0(doc.tables[0], audit_set, team_str, stage1, stage2)

    if len(doc.tables) >= 4:
        members_ctx = _build_committee_context(audit_set)
        _fill_table3_committee(doc.tables[3], members_ctx)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _safe_fill_table0(t0, audit_set, team_str: str, stage1, stage2) -> None:
    """Best-effort fill of FR.233 Table 0 — silently skips rows that don't match."""
    rows = t0.rows
    pairs = [
        (0, 1, audit_set.plan_number or ""),
        (1, 1, audit_set.company_name or ""),
        (2, 1, audit_set.company_address or ""),
        (3, 1, ", ".join(audit_set.standards or [])),
        (4, 1, audit_set.ea_code or ""),
        (6, 1, team_str),
        (7, 1, _fmt_d(stage1.audit_date_start if stage1 else None)),
        (7, 3, _fmt_d(stage2.audit_date_start if stage2 else None)),
        (8, 1, _fmt_d(getattr(stage1, "report_date", None))),
        (8, 3, _fmt_d(getattr(stage2, "report_date", None))),
        (9, 1, _fmt_d(date.today())),
    ]
    for ri, ci, value in pairs:
        if ri < len(rows) and ci < len(rows[ri].cells):
            _set_cell_text(rows[ri].cells[ci]._tc, value)



