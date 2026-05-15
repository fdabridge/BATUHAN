"""
BATUHAN — FR.223 Audit Plan Generator.
Produces a filled audit plan DOCX from AuditPlanInput.
Nothing is persisted — returns raw bytes for streaming download.
"""
from __future__ import annotations
import io
from dataclasses import dataclass, field
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.settings import get_settings


# ── Data contract ──────────────────────────────────────────────────────────────

@dataclass
class AuditPlanInput:
    company_name: str
    company_address: str
    standard_code: str          # e.g. "QMS (ISO 9001:2015)"
    accreditation_body: str     # "UAF" or "TURKAK"
    stage: int                  # 1 or 2
    audit_date: str             # "2026-06-10" or "2026-06-10 / 2026-06-11"
    lead_auditor_name: str
    assignments: list[dict]     # [{auditor_name, role, assigned_clauses:[{clause_id,title}]}]
    opening_time: str = "09:00"
    closing_time: str = "17:00"
    document_ref: str = "FR.223"


# ── Internal helpers ───────────────────────────────────────────────────────────

_DARK_BLUE = "1F4E79"
_LIGHT_GRAY = "F2F2F2"
_TEXT_GRAY  = "808080"


def _parse_time(t: str) -> datetime:
    return datetime.strptime(t, "%H:%M")


def _fmt_slot(start: datetime, end: datetime) -> str:
    return f"{start.strftime('%H:%M')} – {end.strftime('%H:%M')}"


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: str | None = None,
                   size_pt: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_para(doc: Document, text: str, bold: bool = False, size_pt: int = 11,
              align=WD_ALIGN_PARAGRAPH.LEFT, color: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _build_schedule_rows(data: AuditPlanInput) -> list[tuple[str, str, str]]:
    """Return list of (time_slot, activity, auditor_name) tuples."""
    cursor = _parse_time(data.opening_time)
    rows: list[tuple[str, str, str]] = []

    # Opening meeting — 30 min
    opening_end = cursor + timedelta(minutes=30)
    rows.append((_fmt_slot(cursor, opening_end), "Opening Meeting", data.lead_auditor_name))
    cursor = opening_end

    # Clause audit rows — 1 hour per block
    for assignment in data.assignments:
        name = assignment.get("auditor_name", "")
        clauses = assignment.get("assigned_clauses", [])
        if not clauses:
            continue

        # Split into 1-2 blocks depending on clause count
        blocks: list[list[dict]] = []
        if len(clauses) <= 4:
            blocks = [clauses]
        else:
            mid = len(clauses) // 2
            blocks = [clauses[:mid], clauses[mid:]]

        for block in blocks:
            block_end = cursor + timedelta(hours=1)
            activity = ", ".join(
                f"{c['clause_id']} {c['title']}" for c in block
            )
            if len(activity) > 120:
                activity = activity[:117] + "…"
            rows.append((_fmt_slot(cursor, block_end), activity, name))
            cursor = block_end

    # Closing meeting — 30 min
    closing_end = cursor + timedelta(minutes=30)
    rows.append((_fmt_slot(cursor, closing_end), "Closing Meeting", data.lead_auditor_name))

    return rows


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_audit_plan(data: AuditPlanInput) -> bytes:
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # ── Header block ──────────────────────────────────────────────────────────
    _add_para(doc, get_settings().cb_name, bold=True, size_pt=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, "AUDIT PLAN",     bold=True, size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_para(doc, f"Form Ref: {data.document_ref}", size_pt=9,
              align=WD_ALIGN_PARAGRAPH.RIGHT, color=_TEXT_GRAY)
    doc.add_paragraph()

    # ── Info table ────────────────────────────────────────────────────────────
    info_rows = [
        ("Client Organization", data.company_name),
        ("Address",             data.company_address),
        ("Standard",            data.standard_code),
        ("Audit Stage",         f"Stage {data.stage}"),
        ("Audit Date(s)",       data.audit_date),
        ("Accreditation Body",  data.accreditation_body),
        ("Lead Auditor",        data.lead_auditor_name),
    ]
    info_tbl = doc.add_table(rows=len(info_rows), cols=2)
    info_tbl.style = "Table Grid"
    for i, (label, value) in enumerate(info_rows):
        row = info_tbl.rows[i]
        _set_cell_text(row.cells[0], label, bold=True, size_pt=10)
        _set_cell_text(row.cells[1], value, size_pt=10)

    doc.add_paragraph()

    # ── Schedule table ────────────────────────────────────────────────────────
    schedule_rows = _build_schedule_rows(data)
    sched_tbl = doc.add_table(rows=1 + len(schedule_rows), cols=3)
    sched_tbl.style = "Table Grid"

    # Header row
    hdr = sched_tbl.rows[0]
    for cell, txt in zip(hdr.cells, ["Time", "Activity / Clause Reference", "Auditor"]):
        _set_cell_bg(cell, _DARK_BLUE)
        _set_cell_text(cell, txt, bold=True, size_pt=10, color="FFFFFF")

    # Data rows with alternating background
    for i, (slot, activity, auditor) in enumerate(schedule_rows):
        row = sched_tbl.rows[i + 1]
        bg = _LIGHT_GRAY if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        _set_cell_text(row.cells[0], slot,     size_pt=10)
        _set_cell_text(row.cells[1], activity, size_pt=10)
        _set_cell_text(row.cells[2], auditor,  size_pt=10)

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_para(doc, "Prepared by: _____________  Date: _____________", size_pt=10)
    _add_para(doc, "Approved by: _____________  Date: _____________", size_pt=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
