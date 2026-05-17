"""
BATUHAN — Auditor Profile: Document extraction via Claude.

extract_auditor_from_document(file_bytes, filename) -> dict
  Accepts PDF or DOCX bytes, extracts text, then asks Claude to parse
  all auditor profile fields into a structured JSON dict.
  Returns the parsed dict (nulls allowed) or {"error": str}.
"""
from __future__ import annotations
import io
import json
import logging

logger = logging.getLogger(__name__)

_SYSTEM_PROMPT = """\
You are parsing an auditor CV or IFC form (FR.201). Extract all available fields.
Return ONLY valid JSON with these keys:
  name, email, phone, mobile, role,
  education (list of {degree, institution, year}),
  languages (list of {language, level}),
  field_of_expertise,
  ea_codes (list of strings — use ONLY codes from the official IAF EA code list below),
  accreditation_bodies (list of strings like "UAF", "TURKAK"),
  standard_qualifications (list of {standard_code, accreditation_body, scope_category, ea_codes, technical_depth, experience_years}),
  work_experience (list of {employer, position, start_date, end_date, description}),
  training_records (list of {training_date, institution, subject, duration_days, standard_code, certificate_available}).
Use null for any field not found.

OFFICIAL IAF EA CODE LIST — only use codes from this table. Format as "EA N" (e.g. "EA 3"):
EA 1  - Agriculture, forestry and fishing
EA 2  - Mining and quarrying
EA 3  - Food products, beverages and tobacco
EA 4  - Textiles and textile products
EA 5  - Leather and leather products
EA 6  - Wood and wood products
EA 7  - Pulp, paper and paper products
EA 8  - Publishing companies
EA 9  - Printing companies
EA 10 - Manufacture of coke and refined petroleum products
EA 11 - Nuclear fuel
EA 12 - Chemicals, chemical products and fibres
EA 13 - Pharmaceuticals
EA 14 - Rubber and plastic products
EA 15 - Non-metallic mineral products
EA 16 - Concrete, cement, lime, plaster and similar products
EA 17 - Basic metals and fabricated metal products
EA 18 - Machinery and equipment
EA 19 - Electrical and optical equipment
EA 20 - Shipbuilding
EA 21 - Aerospace
EA 22 - Other transport equipment
EA 23 - Manufacturing not elsewhere classified
EA 24 - Recycling
EA 25 - Electricity supply
EA 26 - Gas supply
EA 27 - Water supply
EA 28 - Construction
EA 29 - Wholesale and retail trade; repair of motor vehicles, motorcycles and personal/household goods
EA 30 - Hotels and restaurants
EA 31 - Transport, storage and communication
EA 32 - Financial intermediation, real estate, renting
EA 33 - Information technology
EA 34 - Engineering services
EA 35 - Other services
EA 36 - Public administration
EA 37 - Education
EA 38 - Health and social work
EA 39 - Other social services

ACCREDITATION BODIES RULE:
- accreditation_bodies (top-level field) must ONLY contain the names of national or international
  accreditation bodies — regulatory/governmental bodies that accredit certification bodies.
  Valid examples: "UAF", "TURKAK", "TÜRKAK", "DAkkS", "UKAS", "ANAB", "SCC", "JAS-ANZ",
  "INAB", "COFRAC", "NAB", "RvA", "ILAC", "EA", "SWEDAC", "FINAS".
  NEVER include: company names, employer names, CB (certification body) names, client names,
  or any organization that issues certificates to end clients. If a name looks like a company
  (e.g. "Certification Partner Global FZ LLC", "Bureau Veritas", "SGS", "TÜV SÜD") it is a CB,
  NOT an accreditation body — exclude it entirely from accreditation_bodies.
  If unsure, omit it.

STANDARD QUALIFICATIONS RULES:
- standard_code must be the ISO standard reference, e.g. "ISO 9001", "ISO 14001", "ISO 45001".
- For each qualification, include accreditation_body only if it is a genuine AB (see rule above).
- technical_depth: one of "Lead Auditor", "Team Auditor", "Technical Expert".
- experience_years: total years of documented auditing experience for that standard (integer).
- A qualification must ONLY be included if there is evidence of BOTH: (a) relevant
  training/certification AND (b) documented auditing experience. Training alone → exclude.

EA CODES PER STANDARD (ISO 9001, ISO 14001, ISO 45001, ISO 27001 only):
- Include "ea_codes" as a list of IAF EA codes the auditor has audited in for THAT standard.
- If the CV does not specify sectors per standard, use the auditor's overall ea_codes as default.
- For all other standards: set "ea_codes": [] — do not use EA codes for these.

SCOPE CATEGORY — mandatory for category-based standards. Use the exact mapping below.

━━━ ISO 22000 and FSSC 22000 — Food Chain Categories ━━━
Set "scope_category" to a comma-separated list of applicable codes from this table.
Read the auditor's documented food industry experience and map each sector to its category:

  BIII — Pre-process handling of plant products: cleaning, sorting, cooling, packing of whole
         harvested plant products (grains, fruits, vegetables before further processing).
         Keywords: grain handling, fresh produce packing, crop storage, primary plant handling.

  C0   — Animal primary conversion: slaughterhouses, abattoirs, lairage, evisceration,
         bulk chilling/freezing of animal carcasses.
         Keywords: slaughter, abattoir, animal primary processing.

  CI   — Processing of perishable animal products: meat processing (sausages, deli meats,
         cured meats), dairy (milk, cheese, yogurt, butter, ice cream), fish and seafood
         processing, egg products.
         Keywords: meat processing, dairy, fish processing, seafood, chilled animal products.

  CII  — Processing of perishable plant products: fresh-cut vegetables/fruits, juices,
         chilled plant-based products, minimally processed vegetables.
         Keywords: fresh-cut, fruit juice, vegetable processing, chilled plant.

  CIII — Processing of perishable mixed products (animal + plant): ready meals, sandwiches,
         salads with mixed ingredients, meal kits, chilled prepared foods.
         Keywords: ready meals, mixed perishable, prepared foods, sandwiches.

  CIV  — Processing of ambient stable products: confectionery (chocolate, candy, gum,
         biscuits, cookies, snacks, chips, crackers), canned goods, dried products,
         cereals, flour, rice, pasta, edible oils, sauces, condiments, frozen foods
         (products frozen after processing), beverages (ambient stable: juices in cartons,
         soft drinks, bottled water).
         Keywords: confectionery, chocolate, candy, gum, snacks, biscuits, canned, ambient,
         dried, cereal, edible oil, beverage, frozen.

  D    — Production of feed and pet food: animal feed manufacturing, pet food (dry, wet,
         treats), aquafeed.
         Keywords: animal feed, pet food, livestock feed.

  E    — Catering: food delivered directly to consumers on-site or off-site — restaurants,
         hotel kitchens, school/hospital canteens, catering companies, take-away.
         Keywords: catering, restaurant, hotel kitchen, canteen, food service.

  FI   — Wholesale, retail, e-commerce of food products: supermarkets, food wholesalers,
         online food retail, food distributors who also take physical possession of products.
         Keywords: retail, wholesale, supermarket, food distribution, e-commerce food.

  FII  — Food brokering and trading: food traders, brokers, importers/exporters who never
         take physical possession of the product.
         Keywords: food broker, food trader, food import/export (no physical handling).

  G    — Storage, distribution, transport: cold chain logistics, temperature-controlled
         warehousing, food transport, third-party logistics for food.
         Keywords: cold storage, food logistics, food transport, warehousing, distribution.

  I    — Production of food packaging and packaging materials: manufacturers of food-contact
         packaging (films, containers, cans, bottles, cartons, labels).
         Keywords: food packaging, packaging materials, food-contact materials.

  K    — Production of (bio)chemicals, bio-cultures, food ingredients, processing aids:
         food additives, flavors, colors, preservatives, enzymes, starter cultures,
         cleaning/disinfection agents for food industry, food-grade chemicals.
         Keywords: food additives, flavors, enzymes, food ingredients, processing aids,
         food-grade chemicals, bio-cultures.

Example for an auditor with experience in confectionery, processed meats, beverages, and packaging:
  "scope_category": "CI, CIV, I"
  (CI = processed meats; CIV = confectionery + beverages; I = packaging)

━━━ ISO 13485 — Medical Device Technical Areas ━━━
Set "scope_category" to a comma-separated list from:

  A1.1 — Non-active medical devices (general): bandages, wound care, catheters, surgical
          instruments, syringes, dental devices (non-electronic), diagnostic test strips,
          non-implantable orthopedic devices.
  A1.2 — Non-active implantable: hip/knee replacements, dental implants, bone screws/plates,
          vascular stents, hernia mesh, non-electronic implants.
  A1.3 — Active (non-implantable): diagnostic imaging (X-ray, ultrasound, MRI, CT),
          patient monitoring, surgical energy devices, infusion pumps, ventilators,
          external defibrillators, electrosurgical equipment.
  A1.4 — Active implantable: pacemakers, implantable defibrillators, cochlear implants,
          neurostimulators, implantable drug delivery.
  A1.5 — Sterilization: EtO sterilization, steam/autoclave, radiation, hydrogen peroxide,
          dry heat, aseptic processing, LTSF sterilization.
  A1.6 — Special technologies: software as medical device (SaMD), nanomaterials, devices
          incorporating human blood/tissue derivatives, biologically active coatings.
  A1.7 — Parts, components, raw materials, services: raw material suppliers, contract
          manufacturers, calibration services, testing labs for medical device industry.
  A2.1 — IVD instruments and software: in-vitro diagnostic analyzers, laboratory instruments.
  A2.2 — IVD reagents, calibrators, controls.
  A2.3 — IVD specimen receptacles: blood collection tubes, sample containers.
  A2.4 — Companion diagnostic devices.

━━━ ISO 50001 — Energy Complexity ━━━
Set "scope_category" to one of: "Low", "Medium", "High"
This reflects the complexity of the energy management systems the auditor has audited:
  High = large industrial sites, multiple energy types, many significant energy uses (>6 SEUs)
  Medium = mid-size manufacturing, 2-3 energy types, moderate SEUs
  Low = small facilities, simple energy profile, 1-2 energy types
If the CV does not provide enough information to determine complexity, set scope_category to null.

━━━ ISO 37001 and ISO 37301 — Sector Type ━━━
Set "scope_category" to one of: "Public", "Private", "Third sector/NGO"
  Public = government agencies, public institutions, state-owned enterprises, municipalities
  Private = commercial companies, corporations, banks, manufacturing firms, consulting firms
  Third sector/NGO = non-profit organizations, NGOs, charities, associations, foundations
Default to "Private" if the auditor's documented experience is primarily with commercial
organizations and no public/NGO context is mentioned.

━━━ ISO 9001, ISO 14001, ISO 45001 — Risk/Complexity Category ━━━
Set "scope_category" to one of: "High", "Medium", "Low" (ISO 14001 also allows "Limited")
Based on the typical risk/complexity of sectors the auditor has audited:
  High risk sectors (ISO 9001): food (EA 3), pharma (EA 13), aerospace (EA 21), nuclear (EA 11),
    medical devices (EA 13), construction (EA 28), healthcare (EA 38)
  Medium: most manufacturing, general services
  Low: simple service organizations, low-complexity office environments
For ISO 14001, "Limited" applies to organizations with very minimal environmental aspects.
Default to "Medium" if unclear.
"""


def _extract_text_docx(file_bytes: bytes) -> str:
    import docx  # python-docx
    doc = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _extract_text_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def extract_auditor_from_document(file_bytes: bytes, filename: str) -> dict:
    """
    Parse an auditor document (PDF or DOCX) into a structured profile dict.
    Returns the dict (with null values for missing fields) or {"error": str}.
    """
    try:
        lower = filename.lower()
        if lower.endswith(".docx") or lower.endswith(".doc"):
            text = _extract_text_docx(file_bytes)
        elif lower.endswith(".pdf"):
            text = _extract_text_pdf(file_bytes)
        else:
            return {"error": f"Unsupported file type: {filename}. Use PDF or DOCX."}

        if not text.strip():
            return {"error": "Could not extract any text from the document."}

        from config.settings import get_settings
        import anthropic

        settings = get_settings()
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

        msg = client.messages.create(
            model=settings.claude_model,
            max_tokens=2048,
            system=_SYSTEM_PROMPT,
            messages=[{
                "role": "user",
                "content": f"Document text:\n\n{text[:12000]}",  # cap at ~12k chars
            }],
        )

        raw = msg.content[0].text.strip()

        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()
        if raw.endswith("```"):
            raw = raw[:-3].strip()

        # Use json_repair to handle any malformed JSON from Claude
        from json_repair import repair_json
        result = json.loads(repair_json(raw))

        # Validate and clean EA codes against official IAF list (EA 1–EA 39)
        VALID_EA_NUMBERS = set(range(1, 40))
        raw_ea = result.get("ea_codes") or []
        cleaned_ea = []
        for code in raw_ea:
            if isinstance(code, str):
                normalized = code.strip().upper()
                num_part = normalized[2:].strip() if normalized.startswith("EA") else normalized
                try:
                    num = int(num_part)
                    if num in VALID_EA_NUMBERS:
                        cleaned_ea.append(f"EA {num}")
                    else:
                        logger.warning("[Auditors/Extractor] Dropped invalid EA code: %s", code)
                except ValueError:
                    logger.warning("[Auditors/Extractor] Could not parse EA code: %s", code)
        result["ea_codes"] = cleaned_ea

        # Validate per-qualification ea_codes against official IAF list
        KNOWN_AB_KEYWORDS = {"UAF", "TURKAK", "TÜRKAK", "DAKKS", "UKAS", "ANAB", "SCC",
                             "JAS-ANZ", "INAB", "COFRAC", "NAB", "RVA", "ILAC", "SWEDAC", "FINAS"}
        for q in result.get("standard_qualifications") or []:
            raw_qual_ea = q.get("ea_codes") or []
            cleaned_qual_ea = []
            for code in raw_qual_ea:
                if isinstance(code, str):
                    normalized = code.strip().upper()
                    num_part = normalized[2:].strip() if normalized.startswith("EA") else normalized
                    try:
                        num = int(num_part)
                        if num in VALID_EA_NUMBERS:
                            cleaned_qual_ea.append(f"EA {num}")
                    except ValueError:
                        pass
            q["ea_codes"] = cleaned_qual_ea

            # Strip CB names from per-qualification accreditation_body
            ab = q.get("accreditation_body") or ""
            if ab and not any(kw in ab.upper() for kw in KNOWN_AB_KEYWORDS):
                q["accreditation_body"] = None
                q["_needs_review"] = True

        # Also strip CB names from top-level accreditation_bodies
        raw_abs = result.get("accreditation_bodies") or []
        result["accreditation_bodies"] = [
            ab for ab in raw_abs
            if any(kw in ab.upper() for kw in KNOWN_AB_KEYWORDS)
        ]

        # Flag qualifications missing accreditation_body for human review
        for q in result.get("standard_qualifications") or []:
            if not q.get("accreditation_body"):
                q["_needs_review"] = True

        logger.info("[Auditors/Extractor] Parsed '%s' — name=%s", filename, result.get("name"))
        return result

    except Exception as exc:
        logger.error("[Auditors/Extractor] Failed on '%s': %s", filename, exc, exc_info=True)
        return {"error": str(exc)}
