"""
BATUHAN — Blank Template Parser (T10)
Parses the uploaded .docx audit report template.
Detects all section headings in order, preserves exact titles.
Produces a TemplateMap used by Prompt B and the DOCX assembly engine.
CRITICAL: Template structure must NEVER be altered.

Detection strategy (applied in order, first match wins):
  1. Word Heading 1/2/3 styles on body paragraphs.
  2. Table cells whose first paragraph is ALL CAPS or fully bold (short text).
     This handles templates that use tables instead of Heading styles.
"""

from __future__ import annotations
import logging
import re
from pathlib import Path
from docx import Document
from schemas.models import TemplateMap, TemplateSection

logger = logging.getLogger(__name__)

# Word heading styles that indicate a report section
HEADING_STYLES = {
    "Heading 1", "Heading 2", "Heading 3",
    "heading 1", "heading 2", "heading 3",
}

# Maximum character length a cell's text can have and still be treated as a heading.
_MAX_HEADING_LEN = 150

# Matches template editorial-instruction text that must never be treated as a section heading.
# Kept in sync with assembly/llm_mapper.py _INSTRUCTION_CELL_RE.
_INSTRUCTION_CELL_RE = re.compile(
    r"THESE TARGETS WILL BE USED FOR FOOD"
    r"|IF NO FOOD YOU CAN DELETE"
    r"|DELETE IF NOT APPLICABLE"
    r"|INSERT TEXT HERE"
    r"|\[Name of reviewer[^\]]*\]"
    r"|\[Name of approver[^\]]*\]"
    r"|\[Insert[^\]]+\]"
    r"|\[ADD[^\]]+\]"
    r"|\[YOUR[^\]]+\]",
    re.IGNORECASE,
)


# ---------------------------------------------------------------------------
# Paragraph helpers
# ---------------------------------------------------------------------------

def _is_heading_para(paragraph) -> bool:
    """Return True if the paragraph uses a Word heading style."""
    return paragraph.style.name in HEADING_STYLES


def _get_paragraph_text(paragraph) -> str:
    """Extract full text from a paragraph including runs."""
    return paragraph.text.strip()


def _is_table_heading_para(paragraph) -> bool:
    """
    Return True if a table-cell paragraph looks like a section heading.

    Criteria (either is sufficient):
      - Text is ALL CAPS (only alphabetic characters matter for the check).
      - Every text-carrying run is bold.

    Additionally the text must be non-empty and ≤ _MAX_HEADING_LEN characters
    so that content paragraphs are never mistaken for headings.
    """
    text = paragraph.text.strip()
    if not text or len(text) > _MAX_HEADING_LEN:
        return False

    # ALL CAPS check — at least one alpha char and all alpha chars are uppercase
    alpha_chars = [c for c in text if c.isalpha()]
    if alpha_chars and all(c.isupper() for c in alpha_chars):
        return True

    # Bold check — every run that carries visible text must be bold
    content_runs = [r for r in paragraph.runs if r.text.strip()]
    if content_runs and all(r.bold for r in content_runs):
        return True

    return False


def _collect_placeholder(paragraphs, start_index: int) -> str:
    """
    Collect any non-heading text immediately following a heading paragraph.
    This captures placeholder text like '[Insert content here]'.
    """
    parts: list[str] = []
    i = start_index
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading_para(p):
            break
        text = _get_paragraph_text(p)
        if text:
            parts.append(text)
        i += 1
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Scanning strategies
# ---------------------------------------------------------------------------

def _scan_paragraphs_for_headings(doc) -> list[TemplateSection]:
    """Scan body paragraphs for Heading 1/2/3 style sections."""
    paragraphs = doc.paragraphs
    sections: list[TemplateSection] = []
    order_index = 0
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        if _is_heading_para(p):
            title = _get_paragraph_text(p)
            if title:
                placeholder = _collect_placeholder(paragraphs, i + 1)
                sections.append(TemplateSection(
                    title=title,
                    original_placeholder=placeholder if placeholder else None,
                    order_index=order_index,
                ))
                order_index += 1
        i += 1
    return sections


def _scan_tables_for_headings(doc) -> list[TemplateSection]:
    """
    Scan all tables in the document for cells that look like section headings
    (bold or ALL CAPS short text).

    Iterates tables in document order; within each table iterates rows then
    cells in reading order.  Merged cells are deduplicated via the underlying
    TC element identity so the same heading is never added twice.
    """
    sections: list[TemplateSection] = []
    order_index = 0
    seen_titles: set[str] = set()

    for table in doc.tables:
        seen_tc_ids: set[int] = set()
        for row in table.rows:
            for cell in row.cells:
                # python-docx returns the same Cell object for merged cells;
                # use the underlying lxml element id to deduplicate.
                tc_id = id(cell._tc)
                if tc_id in seen_tc_ids:
                    continue
                seen_tc_ids.add(tc_id)

                for para in cell.paragraphs:
                    if _is_table_heading_para(para):
                        title = para.text.strip()
                        # Skip editorial instructions — they are never real section headings.
                        if title and _INSTRUCTION_CELL_RE.search(title):
                            logger.debug(
                                "Skipping instruction cell as heading: %r", title[:80]
                            )
                            break
                        if title and title not in seen_titles:
                            seen_titles.add(title)
                            sections.append(TemplateSection(
                                title=title,
                                original_placeholder=None,
                                order_index=order_index,
                            ))
                            order_index += 1
                        break  # at most one heading paragraph per cell

    return sections


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_template(template_path: str) -> TemplateMap:
    """
    Parse a .docx template and return a TemplateMap with all sections in order.

    Each TemplateSection contains:
    - title: exact heading text (must be preserved)
    - original_placeholder: any existing placeholder text under the heading
    - order_index: position in the document (0-based)

    Detection order:
      1. Word Heading 1/2/3 styles on body paragraphs.
      2. Fallback: table cells with bold or ALL CAPS short text.

    Raises ValueError if the file cannot be parsed or has no headings.
    """
    path = Path(template_path)
    if not path.exists():
        raise ValueError(f"Template file not found: {template_path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Template must be a .docx file, got: {path.suffix}")

    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"Failed to open template: {e}")

    # Primary: standard Word heading styles
    sections = _scan_paragraphs_for_headings(doc)

    # Fallback: table-based headings (bold / ALL CAPS cell text)
    if not sections:
        logger.info(
            "No Heading-style paragraphs found in '%s'; "
            "scanning tables for bold or ALL CAPS section titles.",
            path.name,
        )
        sections = _scan_tables_for_headings(doc)

    if not sections:
        raise ValueError(
            "No section headings found in template. "
            "Sections must use Word Heading styles (Heading 1, 2, or 3) "
            "or appear as bold or ALL CAPS text in table cells."
        )

    logger.info("Template parsed: %d sections found in '%s'", len(sections), path.name)
    for s in sections:
        logger.debug("  [%d] %s", s.order_index, s.title)

    return TemplateMap(sections=sections, source_path=str(path.resolve()))


def format_sections_for_prompt(template_map: TemplateMap) -> str:
    """
    Format the template section list as a string for injection into Prompt B.
    Lists all section titles in order so Claude knows exactly what to fill.
    """
    lines = ["The report must contain the following sections in this exact order:\n"]
    for s in template_map.sections:
        lines.append(f"{s.order_index + 1}. {s.title}")
    return "\n".join(lines)

