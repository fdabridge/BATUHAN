"""
BATUHAN — Document Text Extraction Pipeline (T8)
Extracts readable text from PDF, DOCX, and TXT files.
Preserves original filenames. Returns list of ParsedDocument objects.
"""

from __future__ import annotations
import logging
from pathlib import Path
from backend.schemas.models import ParsedDocument

logger = logging.getLogger(__name__)


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF using pdfplumber, fallback to PyMuPDF."""
    text_parts: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        if text_parts:
            return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"pdfplumber failed for {path}: {e}. Trying PyMuPDF.")

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PyMuPDF also failed for {path}: {e}")
        return ""


def extract_text_from_docx(path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed for {path}: {e}")
        return ""


def extract_text_from_txt(path: str) -> str:
    """Read plain text file."""
    try:
        return Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        logger.error(f"TXT read failed for {path}: {e}")
        return ""


def extract_text(path: str) -> str:
    """Route to the correct extractor based on file extension."""
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(path)
    elif suffix == ".txt":
        return extract_text_from_txt(path)
    else:
        logger.warning(f"No text extractor for extension '{suffix}' — file: {path}")
        return ""


def parse_documents(file_paths: list[str]) -> list[ParsedDocument]:
    """
    Extract text from all provided file paths.
    Returns a list of ParsedDocument objects with filename, text, and char_count.
    Files that fail extraction are included with empty text and a warning logged.
    """
    results: list[ParsedDocument] = []
    for path in file_paths:
        filename = Path(path).name
        suffix = Path(path).suffix.lower()
        is_image = suffix in (".png", ".jpg", ".jpeg", ".tiff")

        if is_image:
            # Images are handled by the OCR pipeline (T9), skip here
            logger.info(f"Skipping image file for text extraction (OCR handles it): {filename}")
            continue

        text = extract_text(path)
        if not text.strip():
            logger.warning(f"No text extracted from: {filename}")

        results.append(ParsedDocument(
            filename=filename,
            text=text,
            is_ocr_sourced=False,
            char_count=len(text),
        ))

    return results

