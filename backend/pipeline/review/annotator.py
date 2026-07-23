from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from schemas.models import ReviewResult, ReviewFinding, ReviewFindingSeverity
import datetime, logging, io, re

logger = logging.getLogger(__name__)

# Severity → label prefix for comment text
_SEVERITY_LABELS = {
    ReviewFindingSeverity.CRITICAL: "🔴 CRITICAL",
    ReviewFindingSeverity.MAJOR:    "🟠 MAJOR",
    ReviewFindingSeverity.MINOR:    "🟡 MINOR",
    ReviewFindingSeverity.WARNING:  "🔵 WARNING",
    ReviewFindingSeverity.OK:       "✅ OK",
}

_AUTHOR = "BATUHAN Review"
_INITIALS = "BR"


def build_annotated_docx(
    report_path: str,
    review_result: ReviewResult,
) -> bytes:
    """
    Loads the original report DOCX, inserts Word comments for each finding,
    and returns the annotated document as bytes.
    """
    doc = Document(report_path)

    # Filter to non-OK findings only — no need to comment OK clauses
    findings_to_annotate = [
        f for f in review_result.findings
        if f.finding_type.value != "OK"
    ]

    if not findings_to_annotate:
        # No issues — add a summary paragraph at the end
        _append_summary_section(doc, review_result)
        return _doc_to_bytes(doc)

    # Build comments part in document XML
    comments_part = _get_or_create_comments_part(doc)

    comment_id = 1
    annotated_clauses = set()

    for finding in findings_to_annotate:
        # Find the best paragraph to annotate
        target_para = _find_target_paragraph(doc, finding)

        if target_para is not None:
            _insert_comment(
                doc=doc,
                comments_part=comments_part,
                paragraph=target_para,
                comment_id=comment_id,
                finding=finding,
            )
            annotated_clauses.add(finding.clause_id)
            comment_id += 1
        else:
            logger.warning(
                f"Could not find paragraph for clause {finding.clause_id} "
                f"— will include in summary section"
            )

    # Always append a summary section at the end
    _append_summary_section(doc, review_result)

    return _doc_to_bytes(doc)


def _find_target_paragraph(doc: Document, finding: ReviewFinding):
    """
    Finds the best paragraph to annotate for a given finding.
    Strategy:
    1. Search for exact quote text
    2. Search for clause_id pattern (e.g. "4.1", "A.5.1")
    3. Search for clause title keywords
    4. Return None if nothing found
    """
    # Strategy 1: exact quote match
    if finding.quote and len(finding.quote) > 5:
        quote_lower = finding.quote.lower()[:50]
        for para in doc.paragraphs:
            if quote_lower in para.text.lower():
                return para
        # Also search table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if quote_lower in para.text.lower():
                            return para

    # Strategy 2: clause_id match
    clause_patterns = [
        finding.clause_id,
        finding.clause_id.replace(".", "\\."),
    ]
    for para in doc.paragraphs:
        for pattern in clause_patterns:
            if re.search(r'\b' + re.escape(finding.clause_id) + r'\b', para.text):
                return para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if re.search(
                        r'\b' + re.escape(finding.clause_id) + r'\b', para.text
                    ):
                        return para

    # Strategy 3: first keyword from clause title
    if finding.clause_title:
        keywords = [
            w for w in finding.clause_title.split()
            if len(w) > 4 and w.lower() not in
            {"and", "the", "for", "with", "from", "that", "this", "their"}
        ]
        for kw in keywords[:2]:
            kw_lower = kw.lower()
            for para in doc.paragraphs:
                if kw_lower in para.text.lower() and len(para.text) > 20:
                    return para

    return None


def _insert_comment(
    doc: Document,
    comments_part,
    paragraph,
    comment_id: int,
    finding: ReviewFinding,
):
    """
    Inserts a Word comment on the given paragraph.
    Adds commentRangeStart/End and commentReference to paragraph XML,
    and adds the comment definition to comments_part.
    """
    label = _SEVERITY_LABELS.get(finding.severity, "ℹ️ NOTE")
    comment_text = (
        f"{label}\n"
        f"Type: {finding.finding_type.value}\n"
        f"Clause: {finding.clause_id} — {finding.clause_title}\n\n"
        f"{finding.description}\n\n"
        f"Suggestion: {finding.suggestion}"
    )

    # Add comment to comments.xml
    _add_comment_to_part(comments_part, comment_id, comment_text)

    # Add commentRangeStart at beginning of paragraph
    range_start = OxmlElement('w:commentRangeStart')
    range_start.set(qn('w:id'), str(comment_id))
    paragraph._p.insert(0, range_start)

    # Add commentRangeEnd + commentReference after last run
    range_end = OxmlElement('w:commentRangeEnd')
    range_end.set(qn('w:id'), str(comment_id))
    paragraph._p.append(range_end)

    ref_run = OxmlElement('w:r')
    ref_rpr = OxmlElement('w:rPr')
    ref_style = OxmlElement('w:rStyle')
    ref_style.set(qn('w:val'), 'CommentReference')
    ref_rpr.append(ref_style)
    ref_run.append(ref_rpr)

    comment_ref = OxmlElement('w:commentReference')
    comment_ref.set(qn('w:id'), str(comment_id))
    ref_run.append(comment_ref)
    paragraph._p.append(ref_run)


def _add_comment_to_part(comments_part, comment_id: int, comment_text: str):
    """Adds a <w:comment> element to the comments XML part."""
    comment_el = OxmlElement('w:comment')
    comment_el.set(qn('w:id'), str(comment_id))
    comment_el.set(qn('w:author'), _AUTHOR)
    comment_el.set(qn('w:date'), datetime.datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'))
    comment_el.set(qn('w:initials'), _INITIALS)

    for line in comment_text.split('\n'):
        p_el = OxmlElement('w:p')
        ppr = OxmlElement('w:pPr')
        pstyle = OxmlElement('w:pStyle')
        pstyle.set(qn('w:val'), 'CommentText')
        ppr.append(pstyle)
        p_el.append(ppr)

        r_el = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        rstyle = OxmlElement('w:rStyle')
        rstyle.set(qn('w:val'), 'CommentReference')
        rpr.append(rstyle)
        r_el.append(rpr)

        t_el = OxmlElement('w:t')
        t_el.text = line
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r_el.append(t_el)
        p_el.append(r_el)
        comment_el.append(p_el)

    comments_part.append(comment_el)


def _get_or_create_comments_part(doc: Document):
    """
    Gets or creates the word/comments.xml part in the document package.
    Returns the root <w:comments> element.
    """
    from lxml import etree

    # Check if comments part already exists
    part = doc.part
    try:
        comments_part_obj = part.part_related_by(
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        )
        return comments_part_obj._element
    except Exception:
        pass

    # Create new comments part
    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mo="http://schemas.microsoft.com/office/mac/office/2008/main" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:mv="urn:schemas-microsoft-com:mac:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'mc:Ignorable="mv mo w14 wp14">'
        '</w:comments>'
    )

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_part_obj = Part(
        PackURI('/word/comments.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        comments_xml.encode('utf-8'),
        doc.part.package,
    )

    doc.part.relate_to(
        comments_part_obj,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    )

    root = etree.fromstring(comments_xml.encode('utf-8'))
    comments_part_obj._element = root
    return root


def _add_bullet_paragraph(doc: Document):
    """Add a bullet paragraph even when the template omits Word's bullet style."""
    try:
        return doc.add_paragraph(style="List Bullet")
    except KeyError:
        paragraph = doc.add_paragraph()
        paragraph.add_run("\u2022 ")
        return paragraph


def _append_summary_section(doc: Document, review_result: ReviewResult):
    """Appends a structured review summary at the end of the document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_page_break()

    heading = doc.add_heading('BATUHAN REVIEW SUMMARY', level=1)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Standard: {review_result.standard_code}  |  "
        f"Stage: {review_result.stage}  |  "
        f"Accreditation: {review_result.accreditation_body}"
    ).italic = True

    doc.add_paragraph(
        f"Total Issues: {review_result.total_findings}  |  "
        f"Critical: {review_result.critical_count}  |  "
        f"Major: {review_result.major_count}  |  "
        f"Minor: {review_result.minor_count}  |  "
        f"Warnings: {review_result.warning_count}"
    )

    if review_result.overall_assessment:
        doc.add_heading('Overall Assessment', level=2)
        doc.add_paragraph(review_result.overall_assessment)

    non_ok = [f for f in review_result.findings if f.finding_type.value != "OK"]
    if non_ok:
        doc.add_heading('All Findings', level=2)
        for finding in non_ok:
            label = _SEVERITY_LABELS.get(finding.severity, "NOTE")
            p = _add_bullet_paragraph(doc)
            run = p.add_run(
                f"[{label}] {finding.clause_id} — {finding.finding_type.value}: "
            )
            run.bold = True
            p.add_run(finding.description)


def _doc_to_bytes(doc: Document) -> bytes:
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
