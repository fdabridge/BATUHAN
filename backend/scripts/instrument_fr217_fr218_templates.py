"""Instrument the authoritative FR.217 R15 and FR.218 R9 UAF templates.

The source forms are intentionally kept as visually authoritative Word files.
This script only adds docxtpl data handles and invisible signature markers; it
does not rebuild tables, alter page geometry, or change the form wording.

Usage:
    python instrument_fr217_fr218_templates.py fr217 INPUT.docx OUTPUT.docx
    python instrument_fr217_fr218_templates.py fr218 INPUT.docx OUTPUT.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _set_cell(cell, value: str, *, white: bool = False) -> None:
    """Replace a blank/value cell while retaining the cell/table geometry."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(value)
    run.font.name = "Times New Roman"
    run.font.size = Pt(1 if white else 9)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _replace_first_checkbox(cell, expression: str) -> None:
    """Replace the first visible empty-box glyph without rewriting its label."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if "☐" in run.text or "□" in run.text:
                run.text = run.text.replace("☐", expression, 1).replace("□", expression, 1)
                return
    # Legacy Word form-field checkboxes are not exposed as text by python-docx.
    # They are selected after docxtpl rendering by the package post-processor.
    return


def _replace_text_preserving_style(cell, old: str, new: str) -> None:
    """Replace fixed form metadata without changing its existing formatting."""
    for paragraph in cell.paragraphs:
        combined = "".join(run.text for run in paragraph.runs)
        if old not in combined or not paragraph.runs:
            continue
        paragraph.runs[0].text = combined.replace(old, new)
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    raise ValueError(f"Expected form metadata {old!r} was not found")


def _site_expr(index: int, field: str) -> str:
    return f'{{{{ sites[{index}].{field} if sites|length > {index} else "" }}}}'


def instrument_fr217(source: Path, destination: Path) -> None:
    doc = Document(source)
    if len(doc.tables) != 11 or len(doc.tables[8].rows) != 26:
        raise ValueError("Expected authoritative FR.217 R15 structure (11 tables, 26-row MDQMS table)")

    # The supplied filename and approved revision are dated 21.08.2026, while
    # the source footer accidentally says 22.08.2026. Keep every generated R15
    # copy aligned with the approved revision date.
    footer_tables = doc.sections[0].footer.tables
    if len(footer_tables) != 1:
        raise ValueError("Expected authoritative FR.217 R15 footer table")
    _replace_text_preserving_style(
        footer_tables[0].cell(1, 3), "22.08.2026", "21.08.2026",
    )

    # Organization details.
    for row, value in enumerate((
        "{{ company_name }}", "{{ company_address }}", "{{ phone }}",
        "{{ email }}", "{{ website }}",
    )):
        _set_cell(doc.tables[0].cell(row, 1), value)

    # Requested standards and the UAF accreditation boxes beside them.
    standard_rows = (
        (1, 0, "qms_selected"), (2, 0, "ems_selected"),
        (3, 0, "ohsms_selected"), (4, 0, "fsms_selected"),
        (1, 2, "isms_selected"), (2, 2, "enms_selected"),
        (3, 2, "mdqms_selected"), (4, 2, "abms_selected"),
    )
    for row, col, selected in standard_rows:
        _replace_first_checkbox(doc.tables[1].cell(row, col), f'{{{{ "☒" if {selected} else "☐" }}}}')
        _replace_first_checkbox(
            doc.tables[1].cell(row, col + 1),
            f'{{{{ "☒" if {selected} and accreditation_body == "UAF" else "☐" }}}}',
        )

    audit_type_checks = (
        (0, 'audit_type == "initial" and not is_transfer'),
        (1, "is_transfer"),
        (2, 'audit_type == "scope_extension"'),
        (3, 'audit_type == "change_of_address"'),
    )
    for col, condition in audit_type_checks:
        _replace_first_checkbox(doc.tables[2].cell(0, col), f'{{{{ "☒" if {condition} else "☐" }}}}')

    _set_cell(doc.tables[3].cell(0, 1), "{{ scope_en }}")

    personnel = doc.tables[4]
    _set_cell(personnel.cell(0, 1), "{{ total_employees }}")
    _set_cell(personnel.cell(1, 1), "{{ personnel.office_employees or office_employees }}")
    _set_cell(personnel.cell(1, 6), "{{ personnel.repetitive_employees or repetitive_employees }}")
    _set_cell(personnel.cell(2, 1), "{{ personnel.subcontractors or subcontractors }}")
    _set_cell(personnel.cell(2, 6), "{{ personnel.seasonal or 0 }}")
    _set_cell(personnel.cell(3, 1), '{{ "☒ Y / ☐ N" if personnel.shift_same_process else "☐ Y / ☒ N" }}')
    _set_cell(personnel.cell(3, 6), "{{ sites|length }}")
    _set_cell(personnel.cell(4, 2), "{{ personnel.shift_1_count or '' }}")
    _set_cell(personnel.cell(4, 4), "{{ personnel.shift_2_count or '' }}")
    _set_cell(personnel.cell(4, 6), "{{ personnel.shift_3_count or '' }}")

    sites = doc.tables[5]
    for row in range(1, 6):
        index = row - 1
        _set_cell(sites.cell(row, 1), _site_expr(index, "address"))
        _set_cell(sites.cell(row, 2), _site_expr(index, "process"))
        _set_cell(sites.cell(row, 3), _site_expr(index, "employee_count"))

    _set_cell(doc.tables[6].cell(1, 1), "{{ non_applicable_clauses }}")
    _set_cell(doc.tables[7].cell(1, 1), "{{ application_data.fsms_haccp_studies or '' }}")
    _set_cell(
        doc.tables[7].cell(2, 1),
        '{{ "Yes" if application_data.fsms_seasonal_production else "No" }}',
    )

    # Device-class selections are the MDQMS facts currently collected by the
    # portal.  The application stores descriptive labels, while R15 uses the
    # abbreviated MDR class labels printed on the form.
    device_classes = doc.tables[8].cell(5, 1)
    class_conditions = (
        ('Class I', '"Class I (low risk)" in application_data.mdqms_device_classes'),
        # Is/Im/Ir cannot safely be inferred from the current application
        # fields (sterile/measuring/reusable Class I variants are distinct).
        ('Is/Im/Ir', 'false'),
        ('IIa', '"Class IIa (medium risk)" in application_data.mdqms_device_classes'),
        ('IIb', '"Class IIb (medium-high risk)" in application_data.mdqms_device_classes'),
        (
            'III',
            '"Class III (high risk)" in application_data.mdqms_device_classes '
            'or "Active implantable devices" in application_data.mdqms_device_classes',
        ),
    )
    _set_cell(
        device_classes,
        "    ".join(
            f'{{{{ "☒" if {condition} else "☐" }}}} {label}'
            for label, condition in class_conditions
        ),
    )
    _set_cell(
        doc.tables[8].cell(3, 1),
        '{{ "☒ Yes / ☐ No" if "Active implantable devices" in '
        'application_data.mdqms_device_classes else "☐ Yes / ☒ No" }}',
    )

    integration = doc.tables[9]
    integration_keys = (
        (1, 1, "document_management"), (1, 3, "process_approach"),
        (2, 1, "management_review"), (2, 3, "improvement_mechanism"),
        (3, 1, "internal_audit"), (3, 3, "management_support"),
        (4, 1, "policy_objectives"), (4, 3, "risk_based_thinking"),
    )
    for row, col, key in integration_keys:
        _set_cell(integration.cell(row, col), f'{{{{ "☒" if integration_level.get("{key}") else "☐" }}}}')

    _set_cell(doc.tables[10].cell(0, 1), "{{ representative }}")
    _set_cell(doc.tables[10].cell(1, 1), "{{ today }}")
    _set_cell(doc.tables[10].cell(2, 1), "[SIG:CLIENT]", white=True)

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def _scope_expr(scope_var: str, code: str) -> str:
    return f'{{{{ "☒" if "{code}" in {scope_var} else "☐" }}}}'


def instrument_fr218(source: Path, destination: Path) -> None:
    doc = Document(source)
    if len(doc.tables) != 33 or len(doc.tables[28].rows) != 10:
        raise ValueError("Expected authoritative FR.218 R9 structure (33 tables, 10-row duration table)")

    for row, value in enumerate((
        "{{ today }}", "{{ company_name }}", "{{ company_address }}", "{{ standards_str }}",
    )):
        _set_cell(doc.tables[0].cell(row, 1), value)

    personnel = doc.tables[1]
    _set_cell(personnel.cell(0, 1), "{{ personnel.office_employees or office_employees }}")
    _set_cell(personnel.cell(0, 6), "{{ personnel.repetitive_employees or repetitive_employees }}")
    _set_cell(personnel.cell(1, 1), "{{ personnel.subcontractors or subcontractors }}")
    _set_cell(personnel.cell(1, 6), "{{ personnel.seasonal or 0 }}")
    _set_cell(personnel.cell(2, 1), '{{ "☒ Y / ☐ N" if personnel.shift_same_process else "☐ Y / ☒ N" }}')
    _set_cell(personnel.cell(3, 2), "{{ personnel.shift_1_count or '' }}")
    _set_cell(personnel.cell(3, 4), "{{ personnel.shift_2_count or '' }}")
    _set_cell(personnel.cell(3, 6), "{{ personnel.shift_3_count or '' }}")
    _set_cell(personnel.cell(5, 1), "{{ sites|length }}")
    _set_cell(personnel.cell(5, 6), "{{ sites|length }}")
    _set_cell(personnel.cell(6, 1), "{{ total_employees }}")
    _set_cell(personnel.cell(6, 6), "{{ effective_employees }}")

    for row in range(1, 4):
        index = row - 1
        _set_cell(doc.tables[2].cell(row, 0), _site_expr(index, "address"))
        _set_cell(doc.tables[2].cell(row, 2), _site_expr(index, "employee_count"))
        _set_cell(doc.tables[2].cell(row, 3), _site_expr(index, "audit_days"))

    _set_cell(doc.tables[3].cell(0, 1), "{{ scope_en }}")

    # IAF/EA grid: select every applicable code independently for QMS/EMS/OHSMS.
    ea_table = doc.tables[4]
    for row in range(1, len(ea_table.rows)):
        raw_code = ea_table.cell(row, 0).text.strip()
        code = re.sub(r"[^0-9a-z]", "", raw_code.casefold())
        if not code:
            continue
        for col, scope_var in ((2, "qms_scope_codes"), (3, "ems_scope_codes"), (4, "ohsms_scope_codes")):
            if "☐" in ea_table.cell(row, col).text or "□" in ea_table.cell(row, col).text:
                _replace_first_checkbox(ea_table.cell(row, col), _scope_expr(scope_var, code))

    # FSMS category/subcategory rows.
    fsms_table = doc.tables[5]
    for row in range(1, len(fsms_table.rows)):
        cell = fsms_table.cell(row, 1)
        match = re.search(r"(?:☐|□)\s*([A-Z]+)\s*[–-]", cell.text)
        if match:
            _replace_first_checkbox(cell, _scope_expr("fsms_scope_codes", match.group(1).casefold()))

    # MDQMS technical-area grid. The selected scope is stored at A.1.x level;
    # each matching sub-area remains visible and is marked within that group.
    mdqms_table = doc.tables[7]
    for row in range(1, len(mdqms_table.rows)):
        group_text = mdqms_table.cell(row, 0).text
        match = re.search(r"A\.1\.([1-7])", group_text)
        if match:
            cell = mdqms_table.cell(row, 1)
            if "☐" in cell.text or "□" in cell.text:
                _replace_first_checkbox(cell, _scope_expr("mdqms_scope_codes", f"a1{match.group(1)}"))

    # ISMS supports multiple A/B/C/D technical areas. D remains a grouped
    # selection because the application stores the parent technical area.
    isms_table = doc.tables[9]
    handled_d_cells: set[int] = set()
    for row in range(1, len(isms_table.rows)):
        code = isms_table.cell(row, 0).text.strip().casefold()
        cell = isms_table.cell(row, 3)
        if code in {"a", "b", "c"} and ("☐" in cell.text or "□" in cell.text):
            _replace_first_checkbox(cell, _scope_expr("isms_scope_codes", code))
        elif code == "d" and id(cell._tc) not in handled_d_cells:
            handled_d_cells.add(id(cell._tc))
            for paragraph in cell.paragraphs:
                if "☐" in paragraph.text or "□" in paragraph.text:
                    for run in paragraph.runs:
                        if "☐" in run.text or "□" in run.text:
                            run.text = run.text.replace("☐", _scope_expr("isms_scope_codes", "d"), 1)
                            run.text = run.text.replace("□", _scope_expr("isms_scope_codes", "d"), 1)
                            break

    # Reduction/integration/complexity summaries.
    _set_cell(doc.tables[21].cell(0, 1), "{{ reporting_reduction }}")
    _set_cell(doc.tables[21].cell(2, 1), "{{ reporting_reduction }}")
    integration_keys = (
        "document_management", "management_review", "internal_audit", "policy_objectives",
        "process_approach", "improvement_mechanism", "management_support", "risk_based_thinking",
    )
    for row, key in enumerate(integration_keys, 1):
        _set_cell(doc.tables[22].cell(row, 0), f'{{{{ "☒" if integration_level.get("{key}") else "☐" }}}}')

    _set_cell(doc.tables[23].cell(0, 1), "{{ scope_integration_level }}")
    _set_cell(doc.tables[23].cell(0, 3), "{{ integration_reduction }}")
    _set_cell(doc.tables[23].cell(1, 1), "{{ integration_pct }}%")
    _set_cell(doc.tables[23].cell(1, 3), "{{ integration_reduction }}")
    _set_cell(doc.tables[24].cell(0, 1), '{{ man_day_result_qms.category if man_day_result_qms else "" }}')
    _set_cell(doc.tables[24].cell(1, 1), '{{ man_day_result_ems.category if man_day_result_ems else "" }}')
    _set_cell(doc.tables[24].cell(2, 1), '{{ man_day_result_ohsms.category if man_day_result_ohsms else "" }}')
    _set_cell(doc.tables[25].cell(0, 1), "{{ isms_business_score }}")
    _set_cell(doc.tables[25].cell(1, 1), "{{ isms_it_score }}")
    _set_cell(doc.tables[25].cell(0, 3), '{{ man_day_result_isms.category if man_day_result_isms else "" }}')

    fsms = doc.tables[26]
    _set_cell(fsms.cell(2, 1), '{{ man_day_result_fsms.base_ph1 if man_day_result_fsms else "" }}')
    _set_cell(fsms.cell(2, 3), '{{ man_day_result_fsms.haccp_addition if man_day_result_fsms else "" }}')
    _set_cell(fsms.cell(3, 1), '{{ man_day_result_fsms.base_init if man_day_result_fsms else "" }}')
    _set_cell(fsms.cell(3, 3), '{{ man_day_result_fsms.eps if man_day_result_fsms else "" }}')
    _set_cell(fsms.cell(4, 1), '{{ man_day_result_fsms.site_addition if man_day_result_fsms else "" }}')
    _set_cell(fsms.cell(4, 3), '{{ man_day_result_fsms.ad if man_day_result_fsms else "" }}')

    enms = doc.tables[27]
    enms_values = (
        (1, 1, "enms_energy_tj"), (1, 3, "enms_range_ec"), (1, 5, "enms_fec"),
        (2, 1, "enms_energy_types"), (2, 3, "enms_range_et"), (2, 5, "enms_fet"),
        (3, 1, "enms_seu_count"), (3, 3, "enms_range_seu"), (3, 5, "enms_fseu"),
    )
    for row, col, value in enms_values:
        _set_cell(enms.cell(row, col), f'{{{{ {value} if enms_selected else "" }}}}')
    _set_cell(enms.cell(5, 1), 'C = {{ man_day_result.enms_k if enms_selected else "" }}')
    for row, level in ((7, "High"), (8, "Medium"), (9, "Low")):
        _set_cell(enms.cell(row, 6), f'{{{{ "☒" if man_day_result.enms_complexity == "{level}" else "☐" }}}}')

    # Audit-duration rows use the richer per-standard context and remain in the
    # authoritative fixed eight-row grid (unselected standards simply stay blank).
    duration = doc.tables[28]
    standards = ("qms", "ems", "ohsms", "fsms", "isms", "mdqms", "abms", "enms")
    keys = ("ad", "inc_dec_ad", "intg_reduction", "report_reduction", "stage_1", "stage_2", "surv", "rec")
    for row, suffix in enumerate(standards, 1):
        var = f"man_day_result_{suffix}"
        for col, key in enumerate(keys, 1):
            _set_cell(duration.cell(row, col), f'{{{{ {var}.{key} if {var} else "" }}}}')
    totals = (
        "combined_base", "site_addition_total", "integration_reduction", "reporting_reduction",
        "man_day_result.final_ph1", "man_day_result.final_ph2",
        "man_day_result.final_surv1", "man_day_result.final_recert",
    )
    for col, value in enumerate(totals, 1):
        _set_cell(duration.cell(9, col), f"{{{{ {value} or '' }}}}")

    _set_cell(doc.tables[29].cell(0, 1), "{{ recommended_audit_team }}")
    _set_cell(doc.tables[29].cell(1, 1), "{{ decision_committee_chair_name }}")

    # Signature handles are intentionally white so the rendered document shows
    # only the actual signature image, never its routing key.
    _set_cell(doc.tables[32].cell(1, 0), "[SIG:CB_PLANNER]", white=True)
    _set_cell(doc.tables[32].cell(1, 1), "[SIG:CB_REVIEWER]", white=True)
    _set_cell(doc.tables[32].cell(1, 2), "[SIG:CB_CERT_MANAGER]", white=True)

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def main() -> None:
    if len(sys.argv) != 4 or sys.argv[1] not in {"fr217", "fr218"}:
        raise SystemExit(__doc__)
    kind, source, destination = sys.argv[1], Path(sys.argv[2]), Path(sys.argv[3])
    if kind == "fr217":
        instrument_fr217(source, destination)
    else:
        instrument_fr218(source, destination)


if __name__ == "__main__":
    main()
