"""Instrument the FR.232-1 MDQMS Stage 2 report template.

The source form is maintained as a normal Word document. This script adds the
docxtpl information handles and the two audit-report signature anchors without
changing the form's table structure or page setup.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


INFORMATION_HANDLES = {
    (1, 1, 0): (
        "{{ company_name }}\n"
        "{{ company_address }}\n"
        "Tel: {{ phone }} | E-mail: {{ email }}"
    ),
    (2, 1, 0): "{{ scope_en }}",
    (4, 0, 1): "{{ plan_number }}",
    (4, 1, 1): "{{ report_date }}",
    (4, 2, 1): "{{ standards_str }}",
    (4, 3, 1): "{{ lead_auditor_name }}",
    (4, 4, 1): '{{ auditors[0].name if auditors|length > 0 else "" }}',
    (4, 5, 1): '{{ auditors[1].name if auditors|length > 1 else "" }}',
    (4, 6, 1): (
        '{{ technical_experts[0].name '
        'if technical_experts|length > 0 else "" }}'
    ),
    (4, 7, 1): '{{ trainees[0].name if trainees|length > 0 else "" }}',
    (4, 8, 1): '{{ evaluators[0].name if evaluators|length > 0 else "" }}',
    (4, 9, 1): '{{ observers[0].name if observers|length > 0 else "" }}',
    (4, 10, 1): "{{ representative }}",
    (5, 0, 1): "{{ audit_dates }}",
    (5, 0, 3): "{{ audit_days }}",
    (6, 0, 1): (
        '{{ "☑" if is_initial else "☐" }} Initial Certification  '
        '{{ "☑" if is_surveillance else "☐" }} Surveillance  '
        '{{ "☑" if is_recertification else "☐" }} Recertification  '
        '{{ "☑" if is_special else "☐" }} Special'
    ),
    (8, 2, 1): '{{ sites[0].address if sites|length > 0 else "" }}',
    (8, 2, 2): '{{ audit_dates if sites|length > 0 else "" }}',
    (8, 3, 1): '{{ sites[1].address if sites|length > 1 else "" }}',
    (8, 3, 2): '{{ audit_dates if sites|length > 1 else "" }}',
    (8, 4, 1): '{{ sites[2].address if sites|length > 2 else "" }}',
    (8, 4, 2): '{{ audit_dates if sites|length > 2 else "" }}',
    (12, 1, 1): "{{ total_employees }}",
    (12, 2, 1): "{{ subcontractors }}",
    (12, 3, 1): "{{ effective_employees }}",
    (19, 1, 0): "{{ lead_auditor_name }}",
}

SIGNATURE_HANDLES = {
    (19, 3, 0): "[SIG:LEAD_AUDITOR]",
    (19, 3, 1): "[SIG:CB_CERT_MANAGER]",
}


def _set_cell_text(cell, value: str, *, white: bool = False) -> None:
    """Replace cell content while retaining its paragraph/cell formatting."""
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    run = paragraph.add_run(value)
    if white:
        # White keeps the anchor extractable from PDF text while making it
        # invisible in Word/PDF and preventing visible handle artefacts.
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)


def _replace_in_container(container, old: str, new: str) -> None:
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_container(cell, old, new)


def instrument(source: Path, destination: Path) -> None:
    document = Document(source)
    if len(document.tables) != 20:
        raise ValueError(
            f"Expected the 20-table FR.232-1 source form, found {len(document.tables)}"
        )

    for (table_index, row_index, cell_index), value in INFORMATION_HANDLES.items():
        cell = document.tables[table_index].rows[row_index].cells[cell_index]
        _set_cell_text(cell, value)

    for (table_index, row_index, cell_index), value in SIGNATURE_HANDLES.items():
        cell = document.tables[table_index].rows[row_index].cells[cell_index]
        _set_cell_text(cell, value, white=True)

    # The supplied Stage 2 form carried the Stage 1 code in its repeated footer.
    for section in document.sections:
        _replace_in_container(section.header, "FR.231-1", "FR.232-1")
        _replace_in_container(section.footer, "FR.231-1", "FR.232-1")
    _replace_in_container(document, "FR.231-1", "FR.232-1")

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    instrument(args.source, args.destination)


if __name__ == "__main__":
    main()
