#!/usr/bin/env python3
"""
Portal 49a Part 2 — Replace the static blank org-attendee rows in every
FR.225 template with a docxtpl ``{%tr for emp in org_attendees %}`` loop.

Run from the project root:

    python backend/scripts/update_fr225_org_attendee_rows.py
"""
from __future__ import annotations

import copy
import glob
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SEARCH_ROOTS = [
    "uaf_blank_set copy",
    "backend/uaf_blank_set",
]


def _remove_rows(table, start_idx: int, count: int) -> None:
    tbl = table._tbl
    rows = tbl.findall(qn("w:tr"))
    for row in rows[start_idx : start_idx + count]:
        tbl.remove(row)


def _set_cell_text(tc_el, text: str) -> None:
    """Clear all paragraphs in a <w:tc> and write `text` into a single run on
    the first paragraph (preserving the paragraph's pPr formatting)."""
    paragraphs = tc_el.findall(qn("w:p"))
    if not paragraphs:
        return
    for extra in paragraphs[1:]:
        tc_el.remove(extra)
    p = paragraphs[0]
    pPr = p.find(qn("w:pPr"))
    saved_rPr = None
    for r in p.findall(qn("w:r")):
        if saved_rPr is None:
            rPr = r.find(qn("w:rPr"))
            if rPr is not None:
                saved_rPr = copy.deepcopy(rPr)
        p.remove(r)
    if text == "":
        return
    new_r = etree.SubElement(p, qn("w:r"))
    if saved_rPr is not None:
        new_r.append(saved_rPr)
    new_t = etree.SubElement(new_r, qn("w:t"))
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    if pPr is not None and list(p).index(pPr) != 0:
        p.remove(pPr)
        p.insert(0, pPr)


def _set_row_cell_texts(row_el, texts: list[str | None]) -> None:
    """Set each <w:tc>'s text to `texts[i]`. `None` leaves the cell untouched."""
    cells = row_el.findall(".//" + qn("w:tc"))
    for i, tc in enumerate(cells):
        if i >= len(texts) or texts[i] is None:
            continue
        _set_cell_text(tc, texts[i])


def _has_org_attendees_marker(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            if "org_attendees" in cell.text:
                return True
    return False


def process_file(docx_path: str) -> bool:
    path = Path(docx_path)
    doc = Document(docx_path)

    if len(doc.tables) < 3:
        print(f"  SKIP (fewer than 3 tables): {path.name}")
        return False

    table = doc.tables[2]

    if _has_org_attendees_marker(table):
        print(f"  SKIP (already converted): {path.name}")
        return False

    rows = table.rows
    if len(rows) < 8:
        print(f"  SKIP (too few rows in table 2): {path.name}")
        return False

    row6_text = rows[6].cells[0].text.strip().lower()
    if "audit team" not in row6_text and "denetim" not in row6_text:
        print(f"  SKIP (row 6 is not 'Audit Team'): {path.name} → '{row6_text}'")
        return False

    content_row_xml = copy.deepcopy(rows[7]._tr)
    header_row_xml  = copy.deepcopy(rows[1]._tr)

    _remove_rows(table, 2, 4)

    row_c = copy.deepcopy(header_row_xml)
    _set_row_cell_texts(row_c, ["{%tr endfor %}", "", "", ""])

    row_b = copy.deepcopy(content_row_xml)
    _set_row_cell_texts(row_b, [
        "{{ emp.name }}",
        "{{ emp.role }}",
        "[SIG:ORG_OPENING_{{ emp.sig_key }}]",
        "[SIG:ORG_CLOSING_{{ emp.sig_key }}]",
    ])

    row_a = copy.deepcopy(header_row_xml)
    _set_row_cell_texts(row_a, ["{%tr for emp in org_attendees %}", "", "", ""])

    tbl = table._tbl
    current_rows = tbl.findall(qn("w:tr"))
    header_row_element = current_rows[1]
    header_row_element.addnext(row_c)
    header_row_element.addnext(row_b)
    header_row_element.addnext(row_a)

    doc.save(docx_path)
    print(f"  UPDATED: {docx_path}")
    return True


def main() -> None:
    updated = 0
    skipped = 0
    for root in SEARCH_ROOTS:
        pattern = f"{root}/**/FR.225*.docx"
        for f in sorted(glob.glob(pattern, recursive=True)):
            if "~$" in f:
                continue
            if process_file(f):
                updated += 1
            else:
                skipped += 1
    print(f"\nDone. Updated: {updated}, Skipped: {skipped}")


if __name__ == "__main__":
    main()
