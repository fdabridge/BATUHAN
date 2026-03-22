"""
BATUHAN — Integration Tests: Full A→B→C Pipeline (T34)
Runs the complete evidence extraction → report generation → validation pipeline
with mocked Claude API calls. No real network requests, no real prompts.
"""

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

from backend.schemas.models import (
    ISOStandard, AuditStage, ParsedDocument,
    TemplateMap, TemplateSection, StyleGuidance, CorrectionLog,
)
from backend.pipeline.step_a.orchestrator import run_step_a
from backend.pipeline.step_b.orchestrator import run_step_b
from backend.pipeline.step_c.orchestrator import run_step_c


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

MOCK_STEP_A_RESPONSE = """\
## Company Overview
- The organisation operates a Quality Management System aligned with ISO 9001.
- Senior management demonstrates active commitment to continual improvement.

## Scope of Activities
- The scope covers software development and consulting services at a single site.
- All client-facing processes are included within the QMS boundary.

## Documented Information Identified
- A quality manual version 3.1 was reviewed and found to be current.
- Documented procedures for corrective action are established and maintained.

## Key Processes and Functions
- Product delivery follows a defined stage-gate development lifecycle.
- Customer satisfaction is monitored via post-project surveys.

## Evidence of System Implementation
- Internal audit records for Q1–Q4 were reviewed and demonstrate programme effectiveness.
- Management review minutes from two sessions were available and reviewed.

## Audit-Relevant Records
- Training records for all key personnel roles were observed and up to date.
- Nonconformance log entries were reviewed with appropriate corrective actions closed.

## Identified Gaps and Observations
- No formal documented supplier evaluation procedure was found during the audit.
- Customer complaint response timelines require strengthening.
"""

MOCK_STEP_B_RESPONSE = """\
Section Title:
Introduction and Scope

Content:
This Stage 2 audit of the organisation's Quality Management System (QMS) was conducted per
ISO 9001:2015. The scope encompasses software development and consulting services at the
organisation's primary site. The audit team reviewed documented information, records, and
interviewed process owners to assess effective implementation.

---

Section Title:
Key Findings

Content:
The organisation maintains effective documented information including a quality manual and
corrective action procedures. Internal audit records and management review minutes confirm
active system operation. Training records for key personnel were reviewed and found current.
Limited documented evidence was observed for supplier evaluation procedures.

---
"""

MOCK_STEP_C_RESPONSE = """\
## Final Corrected Report

Section Title:
Introduction and Scope

Content:
This Stage 2 audit of the organisation's Quality Management System (QMS) was conducted per
ISO 9001:2015 (Clause 4.3). The scope encompasses software development and consulting services
at the organisation's primary site. The audit team reviewed documented information, records, and
interviewed process owners to assess effective implementation.

---

Section Title:
Key Findings

Content:
The organisation maintains effective documented information including a quality manual and
corrective action procedures per Clause 7.5. Internal audit records and management review
minutes confirm active system operation. Training records for key personnel were reviewed and
found current. Limited documented evidence was observed for supplier evaluation procedures.

---

## List of Corrections Made

- Added ISO 9001:2015 clause references in Introduction and Key Findings sections.

---
"""


def _make_template_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Introduction and Scope", level=1)
    doc.add_paragraph("[Insert introduction here]")
    doc.add_heading("Key Findings", level=1)
    doc.add_paragraph("[Insert findings here]")
    doc.save(str(path))


def _make_corpus(tmp_path: Path) -> list[ParsedDocument]:
    txt = tmp_path / "company.txt"
    txt.write_text(
        "The organisation maintains a quality manual and corrective action procedures. "
        "Internal audits are conducted quarterly. Management reviews occur biannually.",
        encoding="utf-8",
    )
    return [ParsedDocument(filename="company.txt", text=txt.read_text(), char_count=len(txt.read_text()))]


def _make_template_map(tmp_path: Path) -> TemplateMap:
    from backend.parsers.template_parser import parse_template
    p = tmp_path / "template.docx"
    _make_template_docx(p)
    return parse_template(str(p))


def _make_style_guidance() -> StyleGuidance:
    return StyleGuidance(tone_notes=["Use formal audit language."], blocked_company_names=[])


# ---------------------------------------------------------------------------
# T34 — Step A integration
# ---------------------------------------------------------------------------

class TestStepAIntegration:
    def test_step_a_returns_evidence_with_all_sections(self, tmp_path):
        corpus = _make_corpus(tmp_path)
        with patch("backend.pipeline.step_a.orchestrator._call_claude", return_value=MOCK_STEP_A_RESPONSE), \
             patch("backend.pipeline.step_a.orchestrator._load_prompt_a", return_value="Prompt: {document_corpus} {standard} {stage}"), \
             patch("backend.pipeline.step_a.orchestrator.save_text_artifact", return_value=str(tmp_path / "artifact.txt")):
            evidence = run_step_a("job-e2e-1", corpus, ISOStandard.QMS, AuditStage.STAGE_2)

        assert evidence.job_id == "job-e2e-1"
        assert len(evidence.company_overview) >= 1
        assert len(evidence.identified_gaps) >= 1
        assert len(evidence.audit_relevant_records) >= 1

    def test_step_a_marks_weak_evidence_items(self, tmp_path):
        corpus = _make_corpus(tmp_path)
        with patch("backend.pipeline.step_a.orchestrator._call_claude", return_value=MOCK_STEP_A_RESPONSE), \
             patch("backend.pipeline.step_a.orchestrator._load_prompt_a", return_value="Prompt: {document_corpus} {standard} {stage}"), \
             patch("backend.pipeline.step_a.orchestrator.save_text_artifact", return_value=str(tmp_path / "a.txt")):
            evidence = run_step_a("job-e2e-2", corpus, ISOStandard.QMS, AuditStage.STAGE_2)

        all_items = evidence.identified_gaps
        # "No formal documented supplier evaluation procedure was found" → should be weak
        assert any(item.is_weak for item in all_items)


# ---------------------------------------------------------------------------
# T34 — Step B integration
# ---------------------------------------------------------------------------

class TestStepBIntegration:
    def test_step_b_returns_report_with_template_sections(self, tmp_path):
        corpus = _make_corpus(tmp_path)
        template_map = _make_template_map(tmp_path)
        style_guidance = _make_style_guidance()

        with patch("backend.pipeline.step_a.orchestrator._call_claude", return_value=MOCK_STEP_A_RESPONSE), \
             patch("backend.pipeline.step_a.orchestrator._load_prompt_a", return_value="Prompt: {document_corpus} {standard} {stage}"), \
             patch("backend.pipeline.step_a.orchestrator.save_text_artifact", return_value=str(tmp_path / "a.txt")):
            evidence = run_step_a("job-e2e-3", corpus, ISOStandard.QMS, AuditStage.STAGE_2)

        with patch("backend.pipeline.step_b.orchestrator._call_claude", return_value=MOCK_STEP_B_RESPONSE), \
             patch("backend.pipeline.step_b.orchestrator._load_prompt_b", return_value="Prompt B template"), \
             patch("backend.pipeline.step_b.orchestrator.save_text_artifact", return_value=str(tmp_path / "b.txt")):
            report = run_step_b("job-e2e-3", evidence, template_map, style_guidance, ISOStandard.QMS, AuditStage.STAGE_2)

        assert report.job_id == "job-e2e-3"
        assert len(report.sections) == 2
        titles = [s.title for s in report.sections]
        assert "Introduction and Scope" in titles
        assert "Key Findings" in titles


# ---------------------------------------------------------------------------
# T34 — Full A→B→C pipeline integration
# ---------------------------------------------------------------------------

class TestFullPipelineE2E:
    def test_full_pipeline_produces_validated_report_and_correction_log(self, tmp_path):
        corpus = _make_corpus(tmp_path)
        template_map = _make_template_map(tmp_path)
        style_guidance = _make_style_guidance()

        # Step A
        with patch("backend.pipeline.step_a.orchestrator._call_claude", return_value=MOCK_STEP_A_RESPONSE), \
             patch("backend.pipeline.step_a.orchestrator._load_prompt_a", return_value="Prompt: {document_corpus} {standard} {stage}"), \
             patch("backend.pipeline.step_a.orchestrator.save_text_artifact", return_value=str(tmp_path / "a.txt")):
            evidence = run_step_a("job-e2e-full", corpus, ISOStandard.QMS, AuditStage.STAGE_2)

        # Step B
        with patch("backend.pipeline.step_b.orchestrator._call_claude", return_value=MOCK_STEP_B_RESPONSE), \
             patch("backend.pipeline.step_b.orchestrator._load_prompt_b", return_value="Prompt B template"), \
             patch("backend.pipeline.step_b.orchestrator.save_text_artifact", return_value=str(tmp_path / "b.txt")):
            report = run_step_b("job-e2e-full", evidence, template_map, style_guidance, ISOStandard.QMS, AuditStage.STAGE_2)

        # Step C
        with patch("backend.pipeline.step_c.orchestrator._call_claude", return_value=MOCK_STEP_C_RESPONSE), \
             patch("backend.pipeline.step_c.orchestrator._load_prompt_c", return_value="Prompt C template"), \
             patch("backend.pipeline.step_c.orchestrator.save_text_artifact", return_value=str(tmp_path / "c.txt")):
            validated, correction_log = run_step_c("job-e2e-full", report, evidence, template_map, style_guidance)

        # Validate output contracts
        assert validated.job_id == "job-e2e-full"
        assert len(validated.sections) == 2
        assert correction_log.correction_count >= 1
        assert len(correction_log.corrections) == correction_log.correction_count
        # Template integrity: same section titles as template
        output_titles = {s.title for s in validated.sections}
        assert "Introduction and Scope" in output_titles
        assert "Key Findings" in output_titles

    def test_full_pipeline_section_order_preserved(self, tmp_path):
        corpus = _make_corpus(tmp_path)
        template_map = _make_template_map(tmp_path)
        style_guidance = _make_style_guidance()

        with patch("backend.pipeline.step_a.orchestrator._call_claude", return_value=MOCK_STEP_A_RESPONSE), \
             patch("backend.pipeline.step_a.orchestrator._load_prompt_a", return_value="Prompt: {document_corpus} {standard} {stage}"), \
             patch("backend.pipeline.step_a.orchestrator.save_text_artifact", return_value=str(tmp_path / "a.txt")):
            evidence = run_step_a("job-e2e-order", corpus, ISOStandard.QMS, AuditStage.STAGE_2)

        with patch("backend.pipeline.step_b.orchestrator._call_claude", return_value=MOCK_STEP_B_RESPONSE), \
             patch("backend.pipeline.step_b.orchestrator._load_prompt_b", return_value="Prompt B template"), \
             patch("backend.pipeline.step_b.orchestrator.save_text_artifact", return_value=str(tmp_path / "b.txt")):
            report = run_step_b("job-e2e-order", evidence, template_map, style_guidance, ISOStandard.QMS, AuditStage.STAGE_2)

        with patch("backend.pipeline.step_c.orchestrator._call_claude", return_value=MOCK_STEP_C_RESPONSE), \
             patch("backend.pipeline.step_c.orchestrator._load_prompt_c", return_value="Prompt C template"), \
             patch("backend.pipeline.step_c.orchestrator.save_text_artifact", return_value=str(tmp_path / "c.txt")):
            validated, _ = run_step_c("job-e2e-order", report, evidence, template_map, style_guidance)

        # Order indices must be sequential
        indices = [s.order_index for s in validated.sections]
        assert indices == sorted(indices)

