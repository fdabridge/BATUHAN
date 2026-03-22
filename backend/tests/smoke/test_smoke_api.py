"""
BATUHAN — Smoke Tests (T37)
Verifies that the deployed API is reachable and core endpoints behave correctly.

These tests use FastAPI's TestClient (no live network needed for CI), but
can also be pointed at a real deployment via BASE_URL env variable.

Run locally:
    ANTHROPIC_API_KEY=test-key pytest backend/tests/smoke/ -v

Run against staging:
    BASE_URL=http://staging-host:8000 pytest backend/tests/smoke/ -v
"""

from __future__ import annotations
import io
import json
import os
import pytest
from docx import Document

# ---------------------------------------------------------------------------
# Client selection: TestClient (default) or live HTTP
# ---------------------------------------------------------------------------

BASE_URL = os.environ.get("BASE_URL", "")

if BASE_URL:
    import httpx

    @pytest.fixture(scope="session")
    def client():
        with httpx.Client(base_url=BASE_URL, timeout=30) as c:
            yield c
else:
    from fastapi.testclient import TestClient
    from unittest.mock import patch

    @pytest.fixture(scope="session")
    def client():
        # Patch health_checker so TestClient works without Redis
        with patch("backend.monitoring.health_checker.run_health_checks") as mock_hc:
            mock_hc.return_value = {
                "healthy": True,
                "timestamp": "2026-01-01T00:00:00+00:00",
                "version": "1.0.0",
                "checks": {
                    "redis": {"ok": True, "detail": "Redis reachable"},
                    "disk": {"ok": True, "detail": "10000 MB free on storage volume"},
                    "api_key": {"ok": True, "detail": "API key configured"},
                    "stuck_jobs": {"ok": True, "detail": "No stuck jobs", "stuck_jobs": []},
                },
            }
            from backend.main import app
            with TestClient(app) as c:
                yield c


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_bytes(sections: list[str]) -> bytes:
    """Create a minimal DOCX with Heading-1 sections and return as bytes."""
    doc = Document()
    for title in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"[Content for {title}]")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# S1 — Basic liveness
# ---------------------------------------------------------------------------

class TestLiveness:
    def test_root_is_reachable(self, client):
        r = client.get("/")
        assert r.status_code == 200
        data = r.json()
        assert data["system"] == "BATUHAN"
        assert data["status"] == "operational"

    def test_health_returns_ok(self, client):
        r = client.get("/health")
        assert r.status_code == 200
        assert r.json()["status"] == "ok"

    def test_health_detailed_returns_checks(self, client):
        r = client.get("/health/detailed")
        # Accept 200 (healthy) or 503 (degraded but reachable)
        assert r.status_code in (200, 503)
        data = r.json()
        assert "healthy" in data
        assert "checks" in data
        assert "redis" in data["checks"]
        assert "disk" in data["checks"]
        assert "api_key" in data["checks"]
        assert "stuck_jobs" in data["checks"]

    def test_docs_endpoint_available(self, client):
        r = client.get("/docs")
        assert r.status_code == 200


# ---------------------------------------------------------------------------
# S2 — Job creation endpoint
# ---------------------------------------------------------------------------

class TestJobCreation:
    def test_create_job_returns_job_id(self, client):
        """
        Upload valid inputs — the API should accept them and return a job_id.
        Celery queuing is fire-and-forget; failure to connect to Redis is logged
        but does NOT prevent the 200 response (see tasks.py try/except).
        """
        template_bytes = _make_docx_bytes(["Introduction and Scope", "Key Findings"])
        sample_bytes = b"Sample audit report text for style guidance."
        company_bytes = b"Company quality management system documentation."

        r = client.post(
            "/jobs/create",
            data={"standard": "QMS", "stage": "Stage 1"},
            files={
                "company_documents": ("company.txt", company_bytes, "text/plain"),
                "sample_reports": ("sample.txt", sample_bytes, "text/plain"),
                "template": ("template.docx", template_bytes,
                             "application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document"),
            },
        )
        assert r.status_code == 200
        data = r.json()
        assert "job_id" in data
        assert len(data["job_id"]) > 0

    def test_create_job_rejects_invalid_standard(self, client):
        r = client.post(
            "/jobs/create",
            data={"standard": "INVALID", "stage": "Stage 1"},
            files={
                "company_documents": ("c.txt", b"text", "text/plain"),
                "sample_reports": ("s.txt", b"text", "text/plain"),
                "template": ("t.docx", b"fake", "application/octet-stream"),
            },
        )
        assert r.status_code == 422  # Pydantic validation error

    def test_status_endpoint_returns_404_for_unknown_job(self, client):
        import uuid
        unknown_id = f"smoke-test-nonexistent-{uuid.uuid4()}"
        r = client.get(f"/jobs/{unknown_id}/status")
        assert r.status_code == 404

