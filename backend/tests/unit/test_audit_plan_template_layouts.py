"""FR.223 layout compatibility across English and Turkish blank sets."""

from io import BytesIO

from docx import Document

from audit_plan.docx_filler import _find_schedule_table_index, fill_schedule
from audit_plan.schedule_generator import DaySchedule, Slot
from audit_plan.template_reader import read_template


def _turkish_layout_bytes() -> bytes:
    doc = Document()
    header = doc.add_table(rows=3, cols=2)
    header.cell(0, 0).text = "Standard/s"
    header.cell(0, 1).text = "ISO 22000:2018"
    header.cell(1, 0).text = "Audit Type"
    header.cell(1, 1).text = "Stage 2"
    header.cell(2, 0).text = "Audit Date"
    header.cell(2, 1).text = "28.08.2026"

    sites = doc.add_table(rows=2, cols=4)
    for index, text in enumerate(["Saha/lar", "Adres", "Proses/Faaliyet", "Çalışan Sayısı"]):
        sites.cell(0, index).text = text
    for index, text in enumerate(["1", "Ankara", "Food production", "14"]):
        sites.cell(1, index).text = text

    extra = doc.add_table(rows=1, cols=2)
    extra.cell(0, 0).text = "Aşama 1 denetimi sahada mı?"

    team = doc.add_table(rows=3, cols=4)
    for index, text in enumerate(["", "Ad Soyad", "Atandığı Standard/lar", "EA/NACE Kodu"]):
        team.cell(0, index).text = text
    for index, text in enumerate(["Başdenetçi", "Ayşe Yılmaz", "ISO 22000", "CI"]):
        team.cell(1, index).text = text
    for index, text in enumerate(["Teknik Uzman", "Mehmet Kaya", "ISO 22000", "CIII"]):
        team.cell(2, index).text = text

    schedule = doc.add_table(rows=2, cols=5)
    for index, text in enumerate(["Saat", "Standard", "Madde No", "Proses/Faaliyet", "Denetim Ekibi"]):
        schedule.cell(0, index).text = text
    schedule.cell(1, 0).text = "placeholder"

    signature = doc.add_table(rows=1, cols=1)
    signature.cell(0, 0).text = "Signatures"

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_reader_finds_separate_turkish_team_table_and_canonicalizes_roles():
    ctx = read_template(_turkish_layout_bytes())
    assert [(entry.role, entry.name) for entry in ctx.auditors] == [
        ("Lead Auditor", "Ayşe Yılmaz"),
        ("Technical Expert", "Mehmet Kaya"),
    ]
    assert [(site.address, site.process, site.employees) for site in ctx.sites] == [
        ("Ankara", "Food production", "14"),
    ]


def test_filler_locates_turkish_schedule_by_headers_instead_of_fixed_index():
    source = _turkish_layout_bytes()
    doc = Document(BytesIO(source))
    assert _find_schedule_table_index(doc) == 4

    output = fill_schedule(source, [DaySchedule(
        day_number=1,
        date="28.08.2026",
        site="Ankara",
        slots=[Slot(
            time="09.00 – 09.30",
            is_break=False,
            standard="",
            clauses="",
            activity="Opening Meeting",
            auditors="Ayşe Yılmaz (LA)",
        )],
    )])
    filled = Document(BytesIO(output))
    assert "1. Day (28.08.2026)" in filled.tables[4].cell(1, 0).text
    assert filled.tables[2].cell(0, 0).text == "Aşama 1 denetimi sahada mı?"
