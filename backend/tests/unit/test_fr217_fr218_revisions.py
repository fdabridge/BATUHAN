from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from docx.shared import RGBColor

from audit_set.resolver import _find_template


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_ROOT = ROOT / "uaf_blank_set"
GROUPS = ("9-14-45-22-5001", "13485", "27001")


def _stage_1(group: str) -> Path:
    initial = next(
        child for child in (TEMPLATE_ROOT / group).iterdir()
        if child.is_dir() and child.name.strip() == "Initial Certification"
    )
    return next(child for child in initial.iterdir() if child.is_dir() and child.name == "Stage 1")


def _visible_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)


def test_every_production_group_uses_r15_and_r9() -> None:
    for group in GROUPS:
        folder = _stage_1(group)
        fr217 = _find_template(folder, "FR.217")
        fr218 = _find_template(folder, "FR.218")
        assert fr217 is not None and "_R15&21.08.2026.docx" in fr217.name
        assert fr218 is not None and "_R9&21.08.2026.docx" in fr218.name


def test_fr217_signature_handle_is_white() -> None:
    template = _find_template(_stage_1("9-14-45-22-5001"), "FR.217")
    assert template is not None
    doc = Document(template)
    cell = doc.tables[10].cell(2, 1)
    assert cell.text == "[SIG:CLIENT]"
    colors = {
        run.font.color.rgb
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text
    }
    assert colors == {RGBColor(255, 255, 255)}
    assert doc.sections[0].footer.tables[0].cell(1, 3).text == "21.08.2026"


def test_fr217_mdqms_class_translation_uses_application_values() -> None:
    template = _find_template(_stage_1("9-14-45-22-5001"), "FR.217")
    assert template is not None
    doc = Document(template)
    product_classes = doc.tables[8].cell(5, 1).text
    implant_question = doc.tables[8].cell(3, 1).text

    for stored_value in (
        "Class I (low risk)",
        "Class IIa (medium risk)",
        "Class IIb (medium-high risk)",
        "Class III (high risk)",
        "Active implantable devices",
    ):
        assert stored_value in product_classes or stored_value in implant_question


def test_fr218_r9_mapping_and_signature_handles() -> None:
    template = _find_template(_stage_1("9-14-45-22-5001"), "FR.218")
    assert template is not None
    doc = Document(template)
    assert len(doc.tables) == 33
    assert len(doc.tables[28].rows) == 10
    assert "{{ company_name }}" in doc.tables[0].cell(1, 1).text
    assert "{{ company_name }}" in _visible_text(template)

    expected = ("[SIG:CB_PLANNER]", "[SIG:CB_REVIEWER]", "[SIG:CB_CERT_MANAGER]")
    for col, marker in enumerate(expected):
        cell = doc.tables[32].cell(1, col)
        assert cell.text == marker
        colors = {
            run.font.color.rgb
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.text
        }
        assert colors == {RGBColor(255, 255, 255)}


def test_template_discovery_ignores_tmp_and_prefers_highest_revision(tmp_path: Path) -> None:
    for name in (
        "FR.218_Application_Review_Form_R8&09.10.2025.docx",
        "FR.218_Application_Review_Form_R9&21.08.2026.docx",
        "FR.218_Application_Review_Form_R99.docx.tmp.docx",
    ):
        (tmp_path / name).touch()

    selected = _find_template(tmp_path, "FR.218")

    assert selected is not None
    assert re.search(r"_R9&", selected.name)
