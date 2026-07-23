from __future__ import annotations

import io

from docx import Document

from audit_set.fr222_postprocess import ensure_fr222_possible_audit_dates


def test_possible_audit_dates_postprocess_fills_legacy_blank_row():
    doc = Document()
    table = doc.add_table(rows=3, cols=6)
    headings = [
        "Possible Audit Dates",
        "Stage 1",
        "Stage 2",
        "Surveillance 1",
        "Surveillance 2",
        "Recertification",
    ]
    for index, heading in enumerate(headings):
        table.cell(0, index).text = heading
    table.cell(1, 0).text = "Possible Audit Dates"
    table.cell(2, 1).text = "Clauses to be Audited"

    source = io.BytesIO()
    doc.save(source)
    rendered = ensure_fr222_possible_audit_dates(
        source.getvalue(),
        {
            "stage1_dates": "10–11 June 2026",
            "stage2_dates": "14–16 July 2026",
            "surv1_estimated_date": "15/07/2027",
            "surv2_estimated_date": "14/07/2028",
            "recert_estimated_date": "13/07/2029",
        },
    )

    values = [
        cell.text
        for cell in Document(io.BytesIO(rendered)).tables[0].rows[1].cells
    ]
    assert values == [
        "Possible Audit Dates",
        "10–11 June 2026",
        "14–16 July 2026",
        "15/07/2027",
        "14/07/2028",
        "13/07/2029",
    ]


def test_possible_audit_dates_postprocess_preserves_existing_values():
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    for index, heading in enumerate(
        ["Possible Audit Dates", "Stage 1", "Stage 2"],
    ):
        table.cell(0, index).text = heading
    table.cell(1, 1).text = "Manually approved date"

    source = io.BytesIO()
    doc.save(source)
    rendered = ensure_fr222_possible_audit_dates(
        source.getvalue(),
        {
            "stage1_dates": "10–11 June 2026",
            "stage2_dates": "14–16 July 2026",
        },
    )

    values = [
        cell.text
        for cell in Document(io.BytesIO(rendered)).tables[0].rows[1].cells
    ]
    assert values[1] == "Manually approved date"
    assert values[2] == "14–16 July 2026"
