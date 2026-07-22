from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.oxml.ns import qn

from audit_set.committee_slots import (
    expected_committee_sig_keys,
    is_dynamic_committee_member_sig_key,
    planned_committee_chair,
    planned_committee_members,
    planned_committee_slots,
)
from audit_set.fr233_generator import render_fr233_bytes


def _audit_set(committee_members, **overrides):
    defaults = {
        "committee_members": committee_members,
        "stages": [],
        "personnel": {},
        "plan_number": 1652,
        "company_name": "Example Company",
        "company_address": "Example Address",
        "standards": ["QMS"],
        "ea_code": "17",
        "ea_category": None,
        "ea_technical_area": None,
        "required_scope": None,
    }
    defaults.update(overrides)
    return SimpleNamespace(**defaults)


def _run_color_values(cell):
    return [color.get(qn("w:val")) for color in cell._tc.iter(qn("w:color"))]


def test_planned_committee_drives_static_slot_order():
    audit_set = _audit_set([
        {"id": "member-1", "name": "Member One", "role": "member"},
        {"id": "chair-1", "name": "Chair One", "role": "chairperson"},
        {"id": "member-2", "name": "Member Two", "role": "member"},
    ])

    assert [m["id"] for m in planned_committee_members(audit_set)] == [
        "chair-1",
        "member-1",
        "member-2",
    ]
    assert {
        key: member["id"]
        for key, member in planned_committee_slots(audit_set).items()
    } == {
        "COMMITTEE_CHAIR": "chair-1",
        "COMMITTEE_MEMBER_1": "member-1",
        "COMMITTEE_MEMBER_2": "member-2",
    }
    assert planned_committee_chair(audit_set)["id"] == "chair-1"


def test_expected_keys_support_static_and_legacy_dynamic_documents():
    audit_set = _audit_set([
        {"id": "chair-1", "name": "Chair One", "role": "chairperson"},
        {"id": "member-1", "name": "Member One", "role": "member"},
    ])

    assert expected_committee_sig_keys(audit_set) == {
        "COMMITTEE_CHAIR",
        "COMMITTEE_MEMBER_1",
    }
    assert expected_committee_sig_keys(
        audit_set,
        {"COMMITTEE_MEMBER_chair-1", "COMMITTEE_MEMBER_member-1"},
    ) == {
        "COMMITTEE_MEMBER_chair-1",
        "COMMITTEE_MEMBER_member-1",
    }


def test_static_member_rows_are_not_treated_as_dynamic_auditor_keys():
    assert not is_dynamic_committee_member_sig_key("COMMITTEE_MEMBER_1")
    assert not is_dynamic_committee_member_sig_key("COMMITTEE_MEMBER_2")
    assert is_dynamic_committee_member_sig_key("COMMITTEE_MEMBER_auditor-123")


def test_fr233_render_fills_planned_names_codes_and_static_markers():
    audit_set = _audit_set([
        {
            "id": "chair-1",
            "name": "Chair One",
            "role": "chairperson",
            "ea_codes": ["17", "18"],
        },
        {
            "id": "member-1",
            "name": "Member One",
            "role": "member",
            "ea_codes": ["17"],
        },
    ])
    template = next(
        Path("backend/uaf_blank_set").rglob(
            "Initial Certification*/Stage 2/FR.233*.docx",
        )
    )

    rendered = Document(BytesIO(
        render_fr233_bytes(audit_set, db=None, template_path=template),
    ))
    committee_table = rendered.tables[3]

    assert committee_table.cell(1, 1).text == "Chair One"
    assert committee_table.cell(1, 3).text == "17, 18"
    assert committee_table.cell(1, 5).text == "[SIG:COMMITTEE_CHAIR]"
    assert committee_table.cell(2, 1).text == "Member One"
    assert committee_table.cell(2, 5).text == "[SIG:COMMITTEE_MEMBER_1]"
    assert committee_table.cell(3, 1).text == ""
    assert committee_table.cell(3, 5).text == "[SIG:COMMITTEE_MEMBER_2]"
    assert committee_table.cell(6, 4).text == "[SIG:CB_CERT_MANAGER]"
    assert "FFFFFF" in _run_color_values(committee_table.cell(1, 5))
    assert "FFFFFF" in _run_color_values(committee_table.cell(2, 5))
    assert "FFFFFF" in _run_color_values(committee_table.cell(3, 5))
    assert "FFFFFF" in _run_color_values(committee_table.cell(6, 4))


def test_fr233_render_uses_mdqms_required_scope_not_chair_ea_codes():
    audit_set = _audit_set(
        [
            {
                "id": "chair-1",
                "name": "Chair One",
                "role": "chairperson",
                "ea_codes": ["1", "12", "17"],
            },
        ],
        standards=["MDQMS"],
        required_scope={
            "ISO 13485:2016": {
                "type": "medical",
                "codes": ["MD 0101", "MD 0302"],
            },
        },
        ea_technical_area="ISO 13485: MD 0101, MD 0302",
    )
    template = next(
        Path("backend/uaf_blank_set").rglob(
            "Initial Certification*/Stage 2/FR.233*.docx",
        )
    )

    rendered = Document(
        BytesIO(render_fr233_bytes(audit_set, db=None, template_path=template))
    )
    scope_text = rendered.tables[3].cell(1, 3).text

    assert scope_text == "ISO 13485: MD 0101, MD 0302"
    assert "EA 1" not in scope_text
    assert "EA 12" not in scope_text
