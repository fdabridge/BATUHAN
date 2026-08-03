from __future__ import annotations

import ast
from pathlib import Path

from docx import Document
from sqlalchemy import create_engine, text

from audit_set import db_models


ROOT = Path(__file__).resolve().parents[2]
TEMPLATE = (
    ROOT
    / "uaf_blank_set"
    / "9-14-45-22-5001"
    / "Transfer"
    / "FR.250_Transfer_Application_Control_Form_R3&17.07.2026.docx"
)


def test_fr250_has_only_planner_and_transfer_reviewer_signatures():
    document = Document(TEMPLATE)
    rows = [
        "\n".join(cell.text for cell in row.cells)
        for table in document.tables
        for row in table.rows
    ]
    text_content = "\n".join(rows)

    assert "[SIG:CB_PLANNER]" in text_content
    assert "[SIG:TRANSFER_REVIEWER]" in text_content
    assert "[SIG:TRANSFER_CERT_DECISION]" not in text_content
    assert not any(row.startswith("Certification Decision\n") for row in rows)
    assert "The person taking the certification decision shall be different" not in text_content


def test_fr250_remaining_signature_handles_are_white():
    document = Document(TEMPLATE)
    colors = {}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        marker = run.text.strip()
                        if marker.startswith("[SIG:"):
                            colors[marker] = str(run.font.color.rgb)

    assert colors == {
        "[SIG:CB_PLANNER]": "FFFFFF",
        "[SIG:TRANSFER_REVIEWER]": "FFFFFF",
    }


def test_fr250_seeds_only_two_required_signature_slots():
    source = (ROOT / "audit_set" / "documents_router.py").read_text()
    module = ast.parse(source)
    assignment = next(
        node
        for node in module.body
        if isinstance(node, ast.AnnAssign)
        and isinstance(node.target, ast.Name)
        and node.target.id == "DOC_SIG_SLOTS"
    )
    slots = ast.literal_eval(assignment.value)

    assert slots["transfer_review"] == ["cb_planner", "transfer_reviewer"]


def test_legacy_fr250_certification_decision_marker_is_suppressed():
    source = (ROOT / "audit_set" / "viewer_router.py").read_text()

    assert 'doc.document_type == "transfer_review"' in source
    assert 'return {"TRANSFER_CERT_DECISION"}' in source


def test_legacy_unsigned_fr250_committee_slot_is_retired(monkeypatch):
    test_engine = create_engine("sqlite:///:memory:")
    with test_engine.begin() as connection:
        connection.execute(text(
            "CREATE TABLE audit_document_signatures ("
            "id TEXT PRIMARY KEY, document_type TEXT, signer_role_label TEXT, "
            "required BOOLEAN NOT NULL, signed_at TIMESTAMP NULL)"
        ))
        connection.execute(text(
            "INSERT INTO audit_document_signatures VALUES "
            "('unsigned-old', 'transfer_review', 'committee_chair', TRUE, NULL),"
            "('signed-old', 'transfer_review', 'committee_chair', TRUE, '2026-01-01'),"
            "('other-doc', 'fr233', 'committee_chair', TRUE, NULL)"
        ))

    monkeypatch.setattr(db_models, "engine", test_engine)
    db_models._safe_retire_fr250_certification_decision_slots()

    with test_engine.connect() as connection:
        required = dict(connection.execute(text(
            "SELECT id, required FROM audit_document_signatures ORDER BY id"
        )).all())

    assert required == {
        "other-doc": 1,
        "signed-old": 1,
        "unsigned-old": 0,
    }
