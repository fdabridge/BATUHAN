"""
BATUHAN — Unit Tests: Parsers and Schema Validators (T33)
Tests text extraction, template parsing, style extraction, and schema validation.
"""

import pytest
from pathlib import Path
from docx import Document
from pydantic import ValidationError

from schemas.models import (
    ParsedDocument, TemplateMap, TemplateSection,
    ExtractedEvidence, CorrectionLog, JobStatus, JobState, StyleGuidance,
)
from parsers.text_extractor import (
    extract_text_from_txt, extract_text_from_docx, extract_text, parse_documents,
)
from parsers.template_parser import parse_template, format_sections_for_prompt
from parsers.style_extractor import (
    extract_style_from_sample, build_style_guidance, format_style_guidance_for_prompt,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx_with_headings(path: Path, sections: list[tuple[str, str]]) -> None:
    """Create a minimal DOCX with Heading 1 sections."""
    doc = Document()
    for title, body in sections:
        doc.add_heading(title, level=1)
        if body:
            doc.add_paragraph(body)
    doc.save(str(path))


# ---------------------------------------------------------------------------
# T33 — text_extractor
# ---------------------------------------------------------------------------

class TestTextExtractor:
    def test_extract_txt_reads_content(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("Hello audit world", encoding="utf-8")
        assert "Hello audit world" in extract_text_from_txt(str(f))

    def test_extract_docx_reads_paragraphs(self, tmp_path):
        p = tmp_path / "doc.docx"
        doc = Document()
        doc.add_paragraph("Paragraph one.")
        doc.add_paragraph("Paragraph two.")
        doc.save(str(p))
        result = extract_text_from_docx(str(p))
        assert "Paragraph one" in result
        assert "Paragraph two" in result

    def test_extract_text_routes_by_extension_txt(self, tmp_path):
        txt = tmp_path / "file.txt"
        txt.write_text("This is a text file", encoding="utf-8")
        assert "text file" in extract_text(str(txt))

    def test_extract_text_unknown_extension_returns_empty(self, tmp_path):
        f = tmp_path / "data.xyz"
        f.write_bytes(b"binary data")
        assert extract_text(str(f)) == ""

    def test_parse_documents_skips_images(self, tmp_path):
        """Images are handled by OCR pipeline — text extractor must skip them."""
        img = tmp_path / "scan.png"
        img.write_bytes(b"\x89PNG\r\n")
        docs = parse_documents([str(img)])
        assert docs == []

    def test_parse_documents_processes_txt_file(self, tmp_path):
        txt = tmp_path / "company.txt"
        txt.write_text("Company operates a QMS per ISO 9001.", encoding="utf-8")
        docs = parse_documents([str(txt)])
        assert len(docs) == 1
        assert docs[0].filename == "company.txt"
        assert "QMS" in docs[0].text

    def test_parse_documents_missing_file_returns_empty_text(self, tmp_path):
        docs = parse_documents([str(tmp_path / "missing.txt")])
        assert len(docs) == 1
        assert docs[0].text == ""

    def test_parse_documents_sets_char_count(self, tmp_path):
        txt = tmp_path / "x.txt"
        txt.write_text("12345", encoding="utf-8")
        docs = parse_documents([str(txt)])
        assert docs[0].char_count == 5


# ---------------------------------------------------------------------------
# T33 — template_parser
# ---------------------------------------------------------------------------

class TestTemplateParser:
    def test_parse_template_extracts_sections(self, tmp_path):
        p = tmp_path / "template.docx"
        _make_docx_with_headings(p, [
            ("Introduction and Scope", "[Insert]"),
            ("Key Findings", "[Findings]"),
        ])
        tm = parse_template(str(p))
        assert len(tm.sections) == 2
        assert tm.sections[0].title == "Introduction and Scope"
        assert tm.sections[1].title == "Key Findings"

    def test_parse_template_preserves_order_indices(self, tmp_path):
        p = tmp_path / "t.docx"
        _make_docx_with_headings(p, [("A", ""), ("B", ""), ("C", "")])
        tm = parse_template(str(p))
        assert [s.order_index for s in tm.sections] == [0, 1, 2]

    def test_parse_template_sets_source_path(self, tmp_path):
        p = tmp_path / "template.docx"
        _make_docx_with_headings(p, [("Section 1", "")])
        tm = parse_template(str(p))
        assert "template.docx" in tm.source_path

    def test_parse_template_raises_if_no_headings(self, tmp_path):
        p = tmp_path / "empty.docx"
        doc = Document()
        doc.add_paragraph("No headings here — just normal text")
        doc.save(str(p))
        with pytest.raises(ValueError, match="No heading-style sections"):
            parse_template(str(p))

    def test_parse_template_raises_on_missing_file(self, tmp_path):
        with pytest.raises(ValueError, match="not found"):
            parse_template(str(tmp_path / "nonexistent.docx"))

    def test_format_sections_for_prompt_lists_all_titles(self, tmp_path):
        p = tmp_path / "t.docx"
        _make_docx_with_headings(p, [("Introduction", ""), ("Findings", "")])
        tm = parse_template(str(p))
        text = format_sections_for_prompt(tm)
        assert "Introduction" in text
        assert "Findings" in text
        assert "1." in text and "2." in text


# ---------------------------------------------------------------------------
# T33 — style_extractor
# ---------------------------------------------------------------------------

class TestStyleExtractor:
    def test_extract_style_returns_guidance_object(self):
        doc = ParsedDocument(
            filename="sample.pdf",
            text="The audit team confirmed that the organisation demonstrated "
                 "effective implementation. Evidence was observed for corrective action.",
            char_count=100,
        )
        guidance = extract_style_from_sample(doc)
        assert isinstance(guidance.tone_notes, list)
        assert len(guidance.tone_notes) >= 1

    def test_format_style_guidance_includes_blocked_names(self):
        guidance = StyleGuidance(
            tone_notes=["Evidence was observed."],
            structure_notes=[],
            blocked_company_names=["Acme Ltd"],
        )
        text = format_style_guidance_for_prompt(guidance)
        assert "Acme Ltd" in text

    def test_build_style_guidance_returns_merged_object(self, tmp_path):
        txt = tmp_path / "sample.txt"
        txt.write_text(
            "The audit team reviewed records. Evidence was observed. "
            "The organisation demonstrated continual improvement.",
            encoding="utf-8",
        )
        guidance = build_style_guidance([str(txt)])
        assert isinstance(guidance, StyleGuidance)
        assert isinstance(guidance.tone_notes, list)


# ---------------------------------------------------------------------------
# T33 — schema model validation
# ---------------------------------------------------------------------------

class TestSchemaModels:
    def test_template_map_requires_source_path(self):
        with pytest.raises(ValidationError):
            TemplateMap(sections=[])  # source_path missing → should raise

    def test_template_map_valid(self):
        tm = TemplateMap(sections=[], source_path="/tmp/t.docx")
        assert tm.source_path == "/tmp/t.docx"

    def test_extracted_evidence_fields_default_empty(self):
        ev = ExtractedEvidence(job_id="j1", raw_output="")
        assert ev.company_overview == []
        assert ev.identified_gaps == []
        assert ev.audit_relevant_records == []

    def test_correction_log_default_zero_count(self):
        log = CorrectionLog(job_id="j1", corrections=[], correction_count=0)
        assert log.correction_count == 0

    def test_job_status_initial_state(self):
        status = JobStatus(job_id="j1", state=JobState.QUEUED)
        assert status.error_message is None
        assert status.step_timestamps == {}
        assert status.completed_at is None

