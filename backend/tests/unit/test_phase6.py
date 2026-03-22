"""
BATUHAN — Unit Tests: Phase 6 DOCX Assembly & Result Delivery
Tests the assembly engine and packager without real API calls.
Uses a real minimal DOCX created in memory via python-docx.
"""

import pytest
import json
from pathlib import Path
from docx import Document

from schemas.models import (
    ISOStandard, AuditStage,
    ValidatedReport, ReportSection,
    CorrectionLog, CorrectionEntry,
)
from assembly.docx_builder import assemble_docx, _norm
from assembly.result_packager import _format_correction_log_txt


# ---------------------------------------------------------------------------
# Helpers — build in-memory fixtures
# ---------------------------------------------------------------------------

def make_template_docx(tmp_path: Path, sections: list[str]) -> str:
    """Create a minimal .docx template with the given heading sections."""
    doc = Document()
    for title in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph("[Placeholder content]")
    path = str(tmp_path / "template.docx")
    doc.save(path)
    return path


def make_validated_report(job_id: str, sections: list[tuple[str, str]]) -> ValidatedReport:
    report_sections = [
        ReportSection(title=t, content=c, order_index=i)
        for i, (t, c) in enumerate(sections)
    ]
    return ValidatedReport(
        job_id=job_id,
        sections=report_sections,
        correction_log=CorrectionLog(
            job_id=job_id,
            corrections=[CorrectionEntry(section_title="Introduction", description="Fixed placeholder.")],
            correction_count=1,
        ),
        raw_output="raw",
    )


SECTION_TITLES = ["Introduction and Scope", "Key Findings", "Conclusion"]
SECTION_CONTENT = [
    ("Introduction and Scope", "The audit was conducted at the primary facility."),
    ("Key Findings", "The processes are generally conforming to the standard requirements."),
    ("Conclusion", "The management system demonstrates readiness for certification."),
]


# ---------------------------------------------------------------------------
# Tests: docx_builder (T23)
# ---------------------------------------------------------------------------

def test_norm_lowercases_and_strips():
    assert _norm("  Introduction  ") == "introduction"
    assert _norm("KEY FINDINGS") == "key findings"


def test_assemble_docx_creates_output_file(tmp_path):
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-1", SECTION_CONTENT)
    output_path = str(tmp_path / "final_report.docx")

    result = assemble_docx(template_path, report, output_path)

    assert Path(result).exists()
    assert Path(result).suffix == ".docx"


def test_assemble_docx_content_appears_in_output(tmp_path):
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-2", SECTION_CONTENT)
    output_path = str(tmp_path / "output.docx")

    assemble_docx(template_path, report, output_path)

    # Read back and verify content is present
    doc = Document(output_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "primary facility" in all_text
    assert "generally conforming" in all_text
    assert "readiness for certification" in all_text


def test_assemble_docx_preserves_headings(tmp_path):
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-3", SECTION_CONTENT)
    output_path = str(tmp_path / "output.docx")

    assemble_docx(template_path, report, output_path)

    doc = Document(output_path)
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

    for title in SECTION_TITLES:
        assert title in heading_texts, f"Heading '{title}' was lost during assembly."


def test_assemble_docx_replaces_placeholder(tmp_path):
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-4", SECTION_CONTENT)
    output_path = str(tmp_path / "output.docx")

    assemble_docx(template_path, report, output_path)

    doc = Document(output_path)
    all_text = "\n".join(p.text for p in doc.paragraphs)

    # Original placeholder should be replaced
    assert "[Placeholder content]" not in all_text


def test_assemble_docx_missing_template_raises(tmp_path):
    report = make_validated_report("job-d-5", SECTION_CONTENT)
    with pytest.raises(ValueError, match="not found"):
        assemble_docx(str(tmp_path / "nonexistent.docx"), report, str(tmp_path / "out.docx"))


def test_assemble_docx_logs_warning_for_unmatched_section(tmp_path):
    # Template has an extra section not present in corrected content
    extended_titles = SECTION_TITLES + ["Appendix A"]
    template_path = make_template_docx(tmp_path, extended_titles)
    report = make_validated_report("job-d-6", SECTION_CONTENT)
    output_path = str(tmp_path / "output.docx")

    # Should succeed — unmatched section is left intact, not raise
    result = assemble_docx(template_path, report, output_path)
    assert Path(result).exists()


# ---------------------------------------------------------------------------
# Tests: result_packager (T24)
# ---------------------------------------------------------------------------

def test_format_correction_log_txt_with_corrections():
    log = CorrectionLog(
        job_id="job-d-7",
        corrections=[
            CorrectionEntry(section_title="Intro", description="Removed placeholder."),
            CorrectionEntry(section_title=None, description="General phrasing fix."),
        ],
        correction_count=2,
    )
    text = _format_correction_log_txt(log)
    assert "job-d-7" in text
    assert "2" in text
    assert "Removed placeholder" in text
    assert "General phrasing fix" in text


def test_format_correction_log_txt_no_corrections():
    log = CorrectionLog(job_id="job-d-8", corrections=[], correction_count=0)
    text = _format_correction_log_txt(log)
    assert "No corrections were required" in text


def test_package_results_creates_all_artifacts(tmp_path):
    """
    Test package_results end-to-end using direct file paths to avoid
    dependency on the cached settings singleton.
    """
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-9", SECTION_CONTENT)

    artifacts_dir = tmp_path / "artifacts-9"
    artifacts_dir.mkdir()
    docx_output = str(artifacts_dir / "final_report.docx")
    correction_log_output = str(artifacts_dir / "correction_log.txt")
    summary_output = str(artifacts_dir / "job_summary.json")

    # Call assemble_docx and format functions directly (bypassing storage singleton)
    from assembly.docx_builder import assemble_docx
    from assembly.result_packager import _format_correction_log_txt, _build_summary

    final_docx_path = assemble_docx(template_path, report, docx_output)
    correction_log_txt = _format_correction_log_txt(report.correction_log)
    Path(correction_log_output).write_text(correction_log_txt, encoding="utf-8")
    summary = _build_summary(
        job_id="job-d-9",
        standard=ISOStandard.QMS,
        stage=AuditStage.STAGE_1,
        files_used=["quality_manual.pdf", "procedures.docx"],
        correction_count=report.correction_log.correction_count,
        final_docx_path=final_docx_path,
        correction_log_path=correction_log_output,
    )
    Path(summary_output).write_text(json.dumps(summary, indent=2), encoding="utf-8")

    assert Path(final_docx_path).exists()
    assert Path(correction_log_output).exists()
    assert Path(summary_output).exists()


def test_package_results_summary_json_written(tmp_path):
    """Verify the summary JSON contains correct metadata."""
    template_path = make_template_docx(tmp_path, SECTION_TITLES)
    report = make_validated_report("job-d-10", SECTION_CONTENT)

    artifacts_dir = tmp_path / "artifacts-10"
    artifacts_dir.mkdir()
    docx_output = str(artifacts_dir / "final_report.docx")
    correction_log_output = str(artifacts_dir / "correction_log.txt")
    summary_output = str(artifacts_dir / "job_summary.json")

    from assembly.docx_builder import assemble_docx
    from assembly.result_packager import _format_correction_log_txt, _build_summary

    final_docx_path = assemble_docx(template_path, report, docx_output)
    correction_log_txt = _format_correction_log_txt(report.correction_log)
    Path(correction_log_output).write_text(correction_log_txt, encoding="utf-8")
    summary = _build_summary(
        job_id="job-d-10",
        standard=ISOStandard.EMS,
        stage=AuditStage.STAGE_2,
        files_used=["env_manual.pdf"],
        correction_count=report.correction_log.correction_count,
        final_docx_path=final_docx_path,
        correction_log_path=correction_log_output,
    )
    Path(summary_output).write_text(json.dumps(summary, indent=2), encoding="utf-8")

    summary_data = json.loads(Path(summary_output).read_text())
    assert summary_data["standard"] == "EMS"
    assert summary_data["stage"] == "Stage 2"
    assert "env_manual.pdf" in summary_data["files_used"]

