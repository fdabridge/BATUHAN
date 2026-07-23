from docx import Document
from lxml import etree

from pipeline.review.annotator import (
    _add_bullet_paragraph,
    _get_or_create_comments_part,
)


def test_new_comments_part_has_single_well_formed_root():
    document = Document()

    root = _get_or_create_comments_part(document)

    assert etree.QName(root).localname == "comments"
    reparsed = etree.fromstring(etree.tostring(root))
    assert etree.QName(reparsed).localname == "comments"


def test_bullet_paragraph_falls_back_when_template_has_no_list_bullet_style():
    document = Document()
    style = document.styles["List Bullet"]
    style._element.getparent().remove(style._element)

    paragraph = _add_bullet_paragraph(document)
    paragraph.add_run("Finding")

    assert paragraph.text == "\u2022 Finding"
