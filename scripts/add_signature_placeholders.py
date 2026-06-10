#!/usr/bin/env python3
"""
Inject [SIG:PARTY] placeholder text into UAF blank-set DOCX templates.

Run once from the repo root:
    python3 scripts/add_signature_placeholders.py

Each DOCX in backend/uaf_blank_set/ is modified in-place.
Commit all resulting .docx changes as a single binary-update commit.

Purpose
-------
After LibreOffice converts a filled DOCX to PDF, pdfplumber scans the PDF
text stream and locates the bounding box of each [SIG:...] token. That
coordinate drives the signature overlay box in the in-portal viewer. The
marker is 8pt light-gray italic — invisible in print, present in the PDF
content stream.

Placeholder → signer mapping (used by Prompt 23 coordinate extractor):
  [SIG:CB_PLANNER]       CB Planning Officer
  [SIG:CB_REVIEWER]      Committee Reviewer (appointed for this audit set)
  [SIG:CB_CERT_MANAGER]  Certification Manager
  [SIG:LEAD_AUDITOR]     Lead Auditor
  [SIG:CLIENT]           Client Organisation Representative
  [SIG:AUDITOR_MEMBER]   Audit-team member (FR.224 per-person copy)
"""
import glob
import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
PLACEHOLDER_FONT_SIZE = Pt(8)
PLACEHOLDER_COLOR = RGBColor(180, 180, 180)  # light gray — barely visible in print

# Note: the prompt spec says "uaf_blank_set copy" but the actual folder is "uaf_blank_set"
BASE_DIR = Path(__file__).parent.parent / "backend" / "uaf_blank_set"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _append_sig(cell, sig_key: str) -> None:
    """
    Add [SIG:KEY] as a new paragraph at the end of a cell that already has
    text (e.g. "Sign:" or "Signature").  Keeps the existing label intact.
    """
    para = cell.add_paragraph()
    run = para.add_run(f"[SIG:{sig_key}]")
    run.font.size = PLACEHOLDER_FONT_SIZE
    run.font.color.rgb = PLACEHOLDER_COLOR
    run.font.italic = True


def _set_sig(cell, sig_key: str) -> None:
    """
    Set [SIG:KEY] as the content of an empty cell.  Uses the cell's first
    existing paragraph so we don't create a spurious blank line above.
    """
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    run = para.add_run(f"[SIG:{sig_key}]")
    run.font.size = PLACEHOLDER_FONT_SIZE
    run.font.color.rgb = PLACEHOLDER_COLOR
    run.font.italic = True


# ---------------------------------------------------------------------------
# Per-form processors  (all use doc.tables[-1] = the last/signature table)
# ---------------------------------------------------------------------------

def _fr220_fr221(doc: Document) -> None:
    """FR.220 Quotation / FR.221 Agreement — 4R × 2C, Row 3 = Sign row."""
    t = doc.tables[-1]
    _append_sig(t.rows[3].cells[0], "CB_PLANNER")
    _append_sig(t.rows[3].cells[1], "CLIENT")


def _fr222(doc: Document) -> None:
    """FR.222 Audit Programme — 2R × 2C, Row 1 = Signature row."""
    t = doc.tables[-1]
    _append_sig(t.rows[1].cells[0], "CB_PLANNER")
    _append_sig(t.rows[1].cells[1], "CB_CERT_MANAGER")


def _fr223(doc: Document) -> None:
    """FR.223 Audit Plan — 3R × 3C, Row 2 = value row, Col 2 = Signature."""
    t = doc.tables[-1]
    _set_sig(t.rows[2].cells[2], "CLIENT")


def _fr218(doc: Document) -> None:
    """FR.218 Application Review — 3R × 3C, Row 1 = value row (3 signers)."""
    t = doc.tables[-1]
    _set_sig(t.rows[1].cells[0], "CB_PLANNER")
    _set_sig(t.rows[1].cells[1], "CB_REVIEWER")
    _set_sig(t.rows[1].cells[2], "CB_CERT_MANAGER")


def _fr230(doc: Document) -> None:
    """FR.230 NC Form — 1R × 4C, Col 1 = client val, Col 3 = LA val."""
    t = doc.tables[-1]
    _set_sig(t.rows[0].cells[1], "CLIENT")
    _set_sig(t.rows[0].cells[3], "LEAD_AUDITOR")


def _fr231_fr229(doc: Document) -> None:
    """
    FR.231 Stage 1 Report / FR.231-1 MD-QMS Stage 1 Report / FR.229 ISMS Report
    All share: 4R × 2C, Row 3 = actual signature area (below "Signature" label).
    """
    t = doc.tables[-1]
    _set_sig(t.rows[3].cells[0], "LEAD_AUDITOR")
    _set_sig(t.rows[3].cells[1], "CB_REVIEWER")


def _fr232_single(doc: Document) -> None:
    """FR.232_ Audit Report (generic) — 4R × 1C, single reviewer."""
    t = doc.tables[-1]
    _set_sig(t.rows[3].cells[0], "CB_REVIEWER")


def _fr232_double(doc: Document) -> None:
    """FR.232-1_ MD-QMS Audit Report — 4R × 2C, LA + Reviewer."""
    t = doc.tables[-1]
    _set_sig(t.rows[3].cells[0], "LEAD_AUDITOR")
    _set_sig(t.rows[3].cells[1], "CB_REVIEWER")


def _fr211(doc: Document) -> None:
    """FR.211 Auditor Assessment — 2R × 2C, Row 1 Col 1 = Signature."""
    t = doc.tables[-1]
    _set_sig(t.rows[1].cells[1], "CLIENT")


def _fr224(doc: Document) -> None:
    """
    FR.224 Impartiality Declaration — 4R × 3C.
    Row 3 Col 0 = {{ assessed_person_name }} (Jinja2, rendered per auditor).
    Row 3 Col 1 = signature cell.
    """
    t = doc.tables[-1]
    _set_sig(t.rows[3].cells[1], "AUDITOR_MEMBER")


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------

PROCESSORS = [
    # (startswith prefix, handler) — order matters: more specific prefixes first
    ("FR.220", _fr220_fr221),
    ("FR.221", _fr220_fr221),
    ("FR.222", _fr222),
    ("FR.223", _fr223),
    ("FR.218", _fr218),
    ("FR.230", _fr230),
    ("FR.231", _fr231_fr229),   # catches both FR.231_ and FR.231-1_
    ("FR.229", _fr231_fr229),
    ("FR.232-1", _fr232_double), # must come BEFORE FR.232_ check
    ("FR.232_", _fr232_single),
    ("FR.211", _fr211),
    ("FR.224", _fr224),
    # FR.225 and FR.234 are intentionally omitted (handled in Prompt 25)
]


def process_file(path: str) -> str:
    basename = os.path.basename(path)
    for prefix, handler in PROCESSORS:
        if basename.startswith(prefix):
            try:
                doc = Document(path)
                handler(doc)
                doc.save(path)
                return f"  OK    {basename}"
            except Exception as exc:
                return f"  ERROR {basename}: {exc}"
    return f"  SKIP  {basename}"


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    pattern = str(BASE_DIR / "**" / "*.docx")
    files = sorted(glob.glob(pattern, recursive=True))

    if not files:
        print(f"ERROR: No DOCX files found under:\n  {BASE_DIR}")
        print("Make sure you're running from the repo root and the path is correct.")
        sys.exit(1)

    print(f"Processing {len(files)} DOCX files in: {BASE_DIR}\n")
    skipped = 0
    updated = 0
    errors = 0
    for path in files:
        result = process_file(path)
        print(result)
        if result.startswith("  OK"):
            updated += 1
        elif result.startswith("  ERROR"):
            errors += 1
        else:
            skipped += 1

    print(f"\nSummary: {updated} updated, {skipped} skipped, {errors} errors.")
    if errors == 0:
        print("Run `git diff --stat` to verify. Commit all modified .docx files.")
    else:
        print("Fix errors above before committing.")
        sys.exit(1)


if __name__ == "__main__":
    main()
